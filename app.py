# ============================================================
#   NoteSaver Pro — app.py   (Upgraded & Hardened v2.0)
#   Author: NoteSaver Pro Team
#   All imports deduplicated, security hardened, code cleaned
# ============================================================

# ──────────────────────────────────────────────────────────
# 1 ▸ STANDARD LIBRARY
# ──────────────────────────────────────────────────────────
import io
import os
import re
import json
import time
import uuid
import shutil
import logging
import traceback
from io import BytesIO
from random import randint
from datetime import datetime, timedelta, timezone
from functools import wraps

# ──────────────────────────────────────────────────────────
# 2 ▸ FLASK CORE
# ──────────────────────────────────────────────────────────
from flask import (
    Flask, g, render_template, request, redirect,
    url_for, flash, jsonify, send_file, session, current_app
)
from flask_login import login_user, logout_user, login_required, current_user
from flask_wtf.csrf import CSRFProtect, generate_csrf
from werkzeug.security import generate_password_hash, check_password_hash
from werkzeug.exceptions import Forbidden
from werkzeug.utils import secure_filename

# ──────────────────────────────────────────────────────────
# 3 ▸ THIRD-PARTY FLASK EXTENSIONS
# ──────────────────────────────────────────────────────────
from flask_migrate import Migrate
from flask_mail import Mail, Message
from flask_limiter import Limiter
from flask_limiter.util import get_remote_address
from flask_socketio import SocketIO, emit
from flask_session import Session
from itsdangerous import URLSafeTimedSerializer, SignatureExpired, BadSignature

# ──────────────────────────────────────────────────────────
# 4 ▸ DOCUMENT / CONTENT PROCESSING
# ──────────────────────────────────────────────────────────
import bleach
import requests
from bs4 import BeautifulSoup
from user_agents import parse as ua_parse
from dotenv import load_dotenv

# ReportLab (PDF)
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib import colors
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Preformatted, Table, TableStyle
)

# python-docx (Word)
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE

# ──────────────────────────────────────────────────────────
# 5 ▸ LOCAL MODULES
# ──────────────────────────────────────────────────────────
from config import Config
from extensions import db, login_manager, csrf
from models import User, Note, Collaborator, UserSession, Review, SupportTicket, DeletedAccount
from forms import RegistrationForm, LoginForm, NoteForm, RequestResetForm, ResetPasswordForm, RequestUsernameForm
from datetime import datetime
import secrets
import threading

# ──────────────────────────────────────────────────────────
# 6 ▸ LOGGER — must be set up BEFORE any optional imports
# ──────────────────────────────────────────────────────────
if not logging.getLogger().handlers:
    logging.basicConfig(
        level=logging.INFO,
        format='%(asctime)s [%(levelname)s] %(name)s: %(message)s'
    )
logger = logging.getLogger(__name__)

load_dotenv()

# ──────────────────────────────────────────────────────────
# 7 ▸ OPTIONAL MODULES (safe to be missing)
# ──────────────────────────────────────────────────────────

# Razorpay
try:
    import razorpay
    RAZORPAY_AVAILABLE = True
    logger.info("Razorpay module loaded")
except ImportError:
    RAZORPAY_AVAILABLE = False
    logger.warning("Razorpay not available — install with: pip install razorpay")

# Redis
try:
    import redis as _redis_module
    redis_client = _redis_module.Redis(host='localhost', port=6379, db=0, decode_responses=True)
    redis_client.ping()
    REDIS_AVAILABLE = True
    logger.info("Redis connected successfully")
except Exception as _redis_err:
    redis_client = None
    REDIS_AVAILABLE = False
    logger.warning(f"Redis not available ({_redis_err}) — using memory fallback")

# AI API (optional internal module)
try:
    from ai_api import ai_blueprint
    HAS_AI_API = True
    logger.info("AI API module loaded")
except ImportError as _ai_err:
    HAS_AI_API = False
    logger.info(f"AI API module not found ({_ai_err}) — AI features disabled")

# unified_flow_backend (optional)
try:
    from unified_flow_backend import (
        register_unified_flow_routes,
        validate_paper_size,
        get_paper_size_config,
        validate_content,
        ContentOptimizer,
        SmartFlowHandler,
        PAPER_SIZES,
    )
    HAS_UNIFIED_FLOW = True
except ImportError as _uf_err:
    HAS_UNIFIED_FLOW = False
    logger.warning(f"unified_flow_backend not found ({_uf_err}) — using fallback validators")

    # ── Fallback stubs so app doesn't crash ──────────────────
    PAPER_SIZES = {'plain': {}, 'a4': {}, 'letter': {}}

    def validate_paper_size(size):
        return (True, '') if size in PAPER_SIZES else (False, f'Unknown paper size: {size}')

    def get_paper_size_config(size):
        return PAPER_SIZES.get(size, {})

    def validate_content(content):
        return True, ''

    class ContentOptimizer:
        pass

    class SmartFlowHandler:
        pass

    def register_unified_flow_routes(app):
        pass

# ──────────────────────────────────────────────────────────
# 8 ▸ FLASK APP INIT
# ──────────────────────────────────────────────────────────
app = Flask(__name__)
# IMPORTANT: Ensure your Config sets UPLOAD_FOLDER and ALLOWED_EXTENSIONS
# Example: UPLOAD_FOLDER = 'static/profile_pics'
app.config.from_object(Config)
# SECRET_KEY must come from environment / Config — never hardcoded here
if not app.config.get('SECRET_KEY') or app.config['SECRET_KEY'] in ('', 'changeme'):
    raise RuntimeError(
        "❌ SECRET_KEY not set! Add SECRET_KEY=<random-64-char-string> to your .env file"
    )
app.config["SESSION_PERMANENT"] = False
app.config['SESSION_USE_SIGNER'] = True
app.config['SESSION_COOKIE_HTTPONLY'] = True
app.config['SESSION_COOKIE_SAMESITE'] = 'Lax'

SECRET_KEY = os.environ.get("SECRET_KEY")

app.config['ADMIN_EMAILS'] = [app.config.get('MAIL_USERNAME', '')]

# ── Razorpay keys (from .env — never hardcode) ────────────
RAZORPAY_KEY_ID     = os.getenv('RAZORPAY_KEY_ID', '')
RAZORPAY_KEY_SECRET = os.getenv('RAZORPAY_KEY_SECRET', '')

razorpay_client = None
if RAZORPAY_AVAILABLE and RAZORPAY_KEY_ID and RAZORPAY_KEY_SECRET:
    try:
        razorpay_client = razorpay.Client(auth=(RAZORPAY_KEY_ID, RAZORPAY_KEY_SECRET))
        logger.info("Razorpay client initialized")
    except Exception as _rp_err:
        logger.error(f"Razorpay client init failed: {_rp_err}")
elif RAZORPAY_AVAILABLE:
    logger.warning("Razorpay keys missing in .env — payment features disabled")

# ===== SERVER SIDE SESSION =====
# /tmp use karo — Render pe available, restart ke beech bhi rehta hai
app.config["SESSION_PERMANENT"] = False
app.config["SESSION_USE_SIGNER"] = True
app.config["SESSION_KEY_PREFIX"] = "notesaver:"
app.config["SESSION_COOKIE_HTTPONLY"] = True
app.config["SESSION_COOKIE_SAMESITE"] = "Lax"
socketio = SocketIO(app, cors_allowed_origins="*")


# Create upload folder if it doesn't exist
with app.app_context():
    if not os.path.exists(app.config.get('UPLOAD_FOLDER', 'static/profile_pics')):
        os.makedirs(app.config.get('UPLOAD_FOLDER', 'static/profile_pics'))


db.init_app(app)
login_manager.init_app(app)
login_manager.login_view = 'login'

# Flask-Session: /tmp use karo — Render pe available hai aur cross-deploy CSRF issue nahi
app.config["SESSION_TYPE"] = "filesystem"
app.config["SESSION_FILE_DIR"] = "/tmp/flask_sessions"
Session(app)

# Create all database tables on startup (Render pe SQLite ephemeral hai — har deploy pe chahiye)
with app.app_context():
    db.create_all()
    logger.info("Database tables created/verified successfully")


@login_manager.unauthorized_handler
def unauthorized_callback():
    """
    Handle unauthorized access properly for both
    browser and API/AJAX requests.
    """

    # Modern AJAX detection (Flask 2+ compatible)
    is_ajax = (
        request.headers.get('X-Requested-With') == 'XMLHttpRequest'
        or request.accept_mimetypes.best == 'application/json'
        or request.path.startswith('/api/')
    )

    if is_ajax:
        return jsonify({
            "error": "Authentication Required",
            "message": "Your session expired. Please login again."
        }), 401

    # Normal browser request → redirect login
    return redirect(url_for('login'))



migrate = Migrate(app, db)

csrf = CSRFProtect(app)
# Register routes
register_unified_flow_routes(app)



# ===== GLOBAL CONTENT SIZE LIMIT =====
MAX_CONTENT_SIZE = 5_000_000   # 5 million characters (~5MB approx)


@app.errorhandler(403)
def forbidden_handler(e):
    """
    Handle 403 Forbidden errors, specifically CSRF errors, by returning JSON 
    for AJAX requests or rendering a 403 page for regular requests.
    """
    
    # Check if the exception relates to CSRF (Flask-WTF raises a Forbidden 
    # exception with 'CSRF' in its description if a token is missing/invalid)
    is_csrf_error = isinstance(e, Forbidden) and 'CSRF' in str(e).upper()
    
    # Check if the request is an AJAX/JSON request
    is_api_request = (
        request.headers.get('X-Requested-With') == 'XMLHttpRequest' or
        'application/json' in request.headers.get('Accept', '')
    )
    
    if is_csrf_error and is_api_request:
        # Return a clean JSON response for JavaScript to handle
        return jsonify({
            "success": False,
            "error": "Security Token Invalid",
            "message": "Security token missing or expired. Please refresh the page and try saving again."
        }), 403 
        
    # Standard 403 error handler: render a template for non-AJAX/non-CSRF issues
    return render_template('403.html'), 403 # Assuming you have a '403.html' template


mail = Mail(app)


 
# ── 3. ROUTES (existing routes ke saath add karo) ──

@app.route('/debug/admin-check')
@login_required
def debug_admin_check():
    return jsonify({
        'your_email': current_user.email,
        'admin_emails': app.config.get('ADMIN_EMAILS', []),
        'match': current_user.email in app.config.get('ADMIN_EMAILS', [])
    })
 
# ---------- Support Page ----------
@app.route('/support')
@login_required
def support():
    return render_template('support.html')
 
 
# ---------- Submit Ticket + Send Email ----------
@app.route('/api/support/send-email', methods=['POST'])
@login_required
def support_send_email():
    try:
        category = request.form.get('category', '').strip()
        subject  = request.form.get('subject',  '').strip()
        message  = request.form.get('message',  '').strip()
 
        if not all([category, subject, message]):
            return jsonify({'success': False, 'message': 'All fields are required.'}), 400
 
        # Generate unique ticket ref
        ref = 'NSP-' + secrets.token_hex(3).upper()
 
        # Save to database
        ticket = SupportTicket(
            ticket_ref = ref,
            user_id    = current_user.id,
            category   = category,
            subject    = subject,
            message    = message,
            status     = 'open'
        )
        db.session.add(ticket)
        db.session.commit()
 
        # Email to admin
        try:
            admin_msg = Message(
                subject  = f'[{category.upper()}] {subject} — Ticket {ref}',
                sender   = app.config.get('MAIL_DEFAULT_SENDER', 'noreply@notesaverpro.com'),
                recipients = [app.config.get('ADMIN_EMAIL', 'admin@notesaverpro.com')],
                body = f"""New Support Ticket Received
══════════════════════════
Ticket Ref : {ref}
From       : {current_user.username} ({current_user.email})
Category   : {category}
Subject    : {subject}
Date       : {datetime.utcnow().strftime('%d %b %Y, %I:%M %p')} UTC

Message:
{message}

══════════════════════════
Reply at: {request.host_url}admin/support
"""
            )
            _dispatch_email(admin_msg)
        except Exception as mail_err:
            logger.warning(f'Admin email queue failed: {mail_err}')
 
        # Confirmation email to user
        try:
            user_msg = Message(
                subject    = f'We received your request — {ref}',
                sender     = app.config.get('MAIL_DEFAULT_SENDER', 'noreply@notesaverpro.com'),
                recipients = [current_user.email],
                body = f"""Hi {current_user.first_name or current_user.username},
 
Thank you for contacting NoteSaver Pro Support.
 
Your ticket has been created:
  Ticket Ref : {ref}
  Subject    : {subject}
  Category   : {category}
 
We typically reply within 2 hours. You can check your ticket status at:
{request.host_url}support/my-tickets
 
Best regards,
NoteSaver Pro Support Team
"""
            )
            _dispatch_email(user_msg)
        except Exception as mail_err:
            logger.warning(f'User confirmation email queue failed: {mail_err}')
 
        return jsonify({'success': True, 'ticket_ref': ref})
 
    except Exception as e:
        db.session.rollback()
        print(f'Support ticket error: {e}')
        return jsonify({'success': False, 'message': 'Something went wrong. Please try again.'}), 500
 
 
# ---------- User: My Tickets ----------
@app.route('/support/my-tickets')
@login_required
def my_tickets():
    tickets = SupportTicket.query.filter_by(user_id=current_user.id)\
                .order_by(SupportTicket.created_at.desc()).all()
    return render_template('my_tickets.html', tickets=tickets)
 
 
# ══════════════ ADMIN ROUTES ══════════════
 
def admin_required(f):
    @wraps(f)
    def decorated(*args, **kwargs):
        if not current_user.is_authenticated:
            return redirect(url_for('login'))
        admin_emails = app.config.get('ADMIN_EMAILS', [])
        if current_user.email not in admin_emails:
            return "Access Denied", 403
        return f(*args, **kwargs)
    return decorated
 
 
# Admin: All Tickets Dashboard
@app.route('/admin/support')
@login_required
@admin_required
def admin_support():
    status_filter = request.args.get('status', 'all')
    query = SupportTicket.query
    if status_filter != 'all':
        query = query.filter_by(status=status_filter)
    tickets = query.order_by(SupportTicket.created_at.desc()).all()
    counts = {
        'all':         SupportTicket.query.count(),
        'open':        SupportTicket.query.filter_by(status='open').count(),
        'in_progress': SupportTicket.query.filter_by(status='in_progress').count(),
        'resolved':    SupportTicket.query.filter_by(status='resolved').count(),
        'closed':      SupportTicket.query.filter_by(status='closed').count(),
    }
    return render_template('admin_support.html', tickets=tickets,
                           status_filter=status_filter, counts=counts)
 
 
# Admin: View Single Ticket
@app.route('/admin/support/<int:ticket_id>')
@login_required
@admin_required
def admin_ticket_detail(ticket_id):
    ticket = SupportTicket.query.get_or_404(ticket_id)
    # Auto mark as in_progress when admin opens it
    if ticket.status == 'open':
        ticket.status = 'in_progress'
        db.session.commit()
    return render_template('admin_ticket_detail.html', ticket=ticket)
 
 
# Admin: Reply to Ticket
@app.route('/admin/support/<int:ticket_id>/reply', methods=['POST'])
@login_required
@admin_required
def admin_reply_ticket(ticket_id):
    ticket = SupportTicket.query.get_or_404(ticket_id)
    reply_text = request.form.get('reply', '').strip()
    new_status  = request.form.get('status', ticket.status)
 
    if not reply_text:
        return jsonify({'success': False, 'message': 'Reply cannot be empty.'}), 400
 
    ticket.admin_reply    = reply_text
    ticket.admin_reply_at = datetime.utcnow()
    ticket.replied_by     = current_user.email
    ticket.status         = new_status
    db.session.commit()
 
    # Email reply to user
    try:
        reply_msg = Message(
            subject    = f'Re: {ticket.subject} [{ticket.ticket_ref}]',
            sender     = app.config.get('MAIL_DEFAULT_SENDER', 'noreply@notesaverpro.com'),
            recipients = [ticket.user.email],
            body = f"""Hi {ticket.user.first_name or ticket.user.username},
 
We have replied to your support ticket.
 
Ticket  : {ticket.ticket_ref}
Subject : {ticket.subject}
Status  : {new_status.replace('_', ' ').title()}
 
Our Reply:
──────────────────────────
{reply_text}
──────────────────────────
 
If you need further help, reply to this email or visit:
{request.host_url}support
 
Best regards,
NoteSaver Pro Support Team
"""
        )
        _dispatch_email(reply_msg)
    except Exception as e:
        logger.warning(f'Reply email queue failed: {e}')
 
    return jsonify({
        'success': True,
        'message': 'Reply sent successfully!',
        'status':  new_status
    })
 
 
# Admin: Change Status Only
@app.route('/admin/support/<int:ticket_id>/status', methods=['POST'])
@login_required
@admin_required
def admin_update_status(ticket_id):
    ticket     = SupportTicket.query.get_or_404(ticket_id)
    new_status = request.json.get('status')
    valid      = ['open', 'in_progress', 'resolved', 'closed']
    if new_status not in valid:
        return jsonify({'success': False, 'message': 'Invalid status'}), 400
    ticket.status = new_status
    db.session.commit()
    return jsonify({'success': True, 'status': new_status})
 
 
# Admin: Delete Ticket
@app.route('/admin/support/<int:ticket_id>/delete', methods=['POST'])
@login_required
@admin_required
def admin_delete_ticket(ticket_id):
    ticket = SupportTicket.query.get_or_404(ticket_id)
    db.session.delete(ticket)
    db.session.commit()
    return jsonify({'success': True})

# Make CSRF token available in all templates - FIXED VERSION
@app.context_processor
def inject_csrf_token():
    return dict(csrf_token=generate_csrf())


    
@app.route("/editor")
@login_required
def editor():
    csrf_token = generate_csrf()
    return render_template("editor.html", csrf_token=csrf_token)
@app.route('/reviews')
def reviews():
    star_filter = request.args.get('stars', 'all')
    plan_filter = request.args.get('plan',  'all')

    query = Review.query.filter_by(is_approved=True)
    if star_filter.isdigit():
        query = query.filter_by(rating=int(star_filter))
    if plan_filter in ('free', 'premium'):
        query = query.filter_by(plan=plan_filter)

    reviews_list = query.order_by(Review.created_at.desc()).all()

    all_reviews = Review.query.filter_by(is_approved=True).all()
    total       = len(all_reviews)
    avg_rating  = round(sum(r.rating for r in all_reviews) / total, 1) if total else 0
    bar_pct     = {i: round(sum(1 for r in all_reviews if r.rating == i) / total * 100) if total else 0
                   for i in range(1, 6)}

    return render_template(
        'reviews.html',
        reviews=reviews_list,
        total=total,
        avg_rating=avg_rating,
        bar_pct=bar_pct,
        star_filter=star_filter,
        plan_filter=plan_filter,
    )


@app.route('/api/reviews/submit', methods=['POST'])
def api_submit_review():
    data   = request.get_json(silent=True) or {}
    name   = (data.get('name')  or '').strip()[:60]
    title  = (data.get('title') or '').strip()[:100]
    body   = (data.get('body')  or '').strip()[:800]
    plan   = data.get('plan') if data.get('plan') in ('free', 'premium') else None

    try:
        rating = int(data.get('rating', 0))
    except (ValueError, TypeError):
        rating = 0

    if not name or not title or len(body) < 30 or not (1 <= rating <= 5):
        return jsonify(success=False, error='Invalid data'), 400

    # Auto-detect plan from logged-in user if not provided
    if plan is None and current_user.is_authenticated:
        plan = 'premium' if current_user.is_premium_active() else 'free'

    review = Review(
        user_id    = current_user.id if current_user.is_authenticated else None,
        name       = name,
        rating     = rating,
        title      = title,
        body       = body,
        plan       = plan or 'free',
        is_approved= True,
    )
    db.session.add(review)
    db.session.commit()
    return jsonify(success=True, review=review.to_dict())


@app.route('/api/reviews/<int:review_id>/helpful', methods=['POST'])
def api_review_helpful(review_id):
    review = db.session.get(Review, review_id)
    if not review:
        return jsonify(success=False, error='Not found'), 404
    review.helpful += 1
    db.session.commit()
    return jsonify(success=True, helpful=review.helpful)


@app.route('/api/reviews/<int:review_id>/delete', methods=['DELETE'])
@login_required
def api_review_delete(review_id):
    review = db.session.get(Review, review_id)
    if not review:
        return jsonify(success=False, error='Review nahi mila'), 404
    if review.user_id != current_user.id:
        return jsonify(success=False, error='Permission nahi hai'), 403
    db.session.delete(review)
    db.session.commit()
    return jsonify(success=True)


# Serializer for generating tokens for password reset
# Serializer uses SECRET_KEY already set from Config / .env
serializer = URLSafeTimedSerializer(app.config['SECRET_KEY'])

@login_manager.user_loader
def load_user(user_id):
    return db.session.get(User, int(user_id))


# ========================================================================
# RAZORPAY PAYMENT INTEGRATION ROUTES
# ========================================================================

@app.route('/api/create-order', methods=['POST'])
@login_required
def create_order():
    """Create Razorpay order for premium subscription"""
    
    if not RAZORPAY_AVAILABLE or not razorpay_client:
        return jsonify({
            "success": False,
            "error": "Payment gateway not configured"
        }), 500
    
    try:
        # Amount in paise (₹99 = 9900 paise)
        amount = 9900
        
        # Create Razorpay order
        order = razorpay_client.order.create({
            "amount": amount,
            "currency": "INR",
            "payment_capture": 1,  # Auto capture payment
            "notes": {
                "user_id": current_user.id,
                "username": current_user.username,
                "email": current_user.email
            }
        })
        
        logger.info(f"Razorpay order created for user {current_user.username}: {order['id']}")
        
        return jsonify({
            "success": True,
            "order_id": order['id'],
            "key": RAZORPAY_KEY_ID,
            "amount": amount,
            "currency": "INR",
            "name": current_user.username,
            "email": current_user.email
        })
        
    except Exception as e:
        logger.error(f"Error creating Razorpay order: {e}")
        return jsonify({
            "success": False,
            "error": str(e)
        }), 500


@app.route('/api/verify-payment', methods=['POST'])
@login_required
def verify_payment():
    """Verify Razorpay payment and activate premium"""
    
    if not RAZORPAY_AVAILABLE or not razorpay_client:
        return jsonify({
            "success": False,
            "error": "Payment gateway not configured"
        }), 500
    
    try:
        data = request.json
        
        # Get payment details from request
        payment_id = data.get('razorpay_payment_id')
        order_id = data.get('razorpay_order_id')
        signature = data.get('razorpay_signature')
        
        if not all([payment_id, order_id, signature]):
            return jsonify({
                "success": False,
                "error": "Missing payment details"
            }), 400
        
        # Verify payment signature
        try:
            razorpay_client.utility.verify_payment_signature({
                'razorpay_order_id': order_id,
                'razorpay_payment_id': payment_id,
                'razorpay_signature': signature
            })
        except razorpay.errors.SignatureVerificationError:
            logger.warning(f"Payment signature verification failed for user {current_user.username}")
            return jsonify({
                "success": False,
                "error": "Payment verification failed"
            }), 400
        
        # Payment verified successfully - Activate Premium
        current_user.is_premium = True
        current_user.premium_expiry = datetime.now(timezone.utc) + timedelta(days=30)
        
        # Store payment record (optional - add Payment model if needed)
        db.session.commit()
        
        logger.info(f"✅ Premium activated for user {current_user.username} - Payment ID: {payment_id}")
        
        return jsonify({
            "success": True,
            "message": "Premium activated successfully!",
            "expiry_date": current_user.premium_expiry.strftime('%d %B %Y')
        })
        
    except Exception as e:
        db.session.rollback()
        logger.error(f"Error verifying payment: {e}")
        return jsonify({
            "success": False,
            "error": "Payment verification failed"
        }), 500


@app.route('/api/payment-status/<order_id>')
@login_required
def payment_status(order_id):
    """Check payment status (optional endpoint)"""
    
    if not RAZORPAY_AVAILABLE or not razorpay_client:
        return jsonify({
            "success": False,
            "error": "Payment gateway not configured"
        }), 500
    
    try:
        order = razorpay_client.order.fetch(order_id)
        
        return jsonify({
            "success": True,
            "order_id": order['id'],
            "status": order['status'],
            "amount": order['amount'],
            "currency": order['currency']
        })
        
    except Exception as e:
        logger.error(f"Error fetching payment status: {e}")
        return jsonify({
            "success": False,
            "error": str(e)
        }), 500


def check_collaboration_permission(view_function):
    """
    Check if current user has permission to access this page
    Based on Collaborator model
    """
    @wraps(view_function)
    def decorated_function(*args, **kwargs):
        # Authentication check
        if not current_user.is_authenticated:
            return redirect(url_for('login'))

        # Get note_id from kwargs
        note_id = kwargs.get('note_id')
        if not note_id:
            return redirect(url_for('dashboard'))

        note = Note.query.get(note_id)
        if not note:
            flash('❌ Note not found', 'danger')
            return redirect(url_for('dashboard'))

        # Owner - full access
        if note.user_id == current_user.id:
            g.user_permission = 'owner'
            g.can_edit = True
            g.can_comment = True
            g.can_delete = True
            g.can_manage = True
            g.note_id = note_id
            g.note = note
            return view_function(*args, **kwargs)

        # Check collaborator record
        collaborator = Collaborator.query.filter_by(
            note_id=note_id,
            collaborator_id=current_user.id,
            status='accepted'
        ).first()

        if not collaborator:
            flash('❌ You do not have access to this note', 'danger')
            logger.warning(f"⚠️ Access denied: User {current_user.id} tried to access note {note_id}")
            return redirect(url_for('dashboard'))

        if collaborator.is_expired():
            flash('❌ Your invitation has expired', 'danger')
            return redirect(url_for('dashboard'))

        g.user_permission = collaborator.permission
        g.can_edit = collaborator.can_edit()
        g.can_comment = collaborator.can_comment()
        g.can_delete = collaborator.can_delete()
        g.can_manage = collaborator.can_manage_collaborators()
        g.note_id = note_id
        g.note = note
        g.note_collaborator = collaborator

        logger.info(f"✅ Access granted: User {current_user.id} ({collaborator.permission}) accessing note {note_id}")

        return view_function(*args, **kwargs)

    return decorated_function
def _dispatch_email(msg):
    """Send email in background thread with app context"""
    try:
        app_obj = app._get_current_object() if hasattr(app, "_get_current_object") else app

        def send_async():
            with app_obj.app_context():
                mail.send(msg)

        thread = threading.Thread(target=send_async)
        thread.start()

    except Exception as e:
        logger.error(f"Email dispatch error: {e}")

def require_permission(required_permission):
    """
    Decorator to require a minimum permission level for a route.
    """
    def decorator(view_function):
        @wraps(view_function)
        def decorated_function(*args, **kwargs):
            if not current_user.is_authenticated:
                return redirect(url_for('login'))

            note_id = kwargs.get('note_id')
            if not note_id:
                return redirect(url_for('dashboard'))

            note = Note.query.get(note_id)
            if not note:
                flash('❌ Note not found', 'danger')
                return redirect(url_for('dashboard'))

            if note.user_id == current_user.id:
                # Owner always allowed
                g.user_permission = 'owner'
                g.can_edit = True
                g.can_comment = True
                g.can_delete = True
                g.can_manage = True
                return view_function(*args, **kwargs)

            collaborator = Collaborator.query.filter_by(
                note_id=note_id,
                collaborator_id=current_user.id,
                status='accepted'
            ).first()

            if not collaborator:
                flash('❌ You do not have access to this note', 'danger')
                return redirect(url_for('dashboard'))

            if not collaborator.has_permission(required_permission):
                logger.warning(f"⚠️ Insufficient permission: User {current_user.id} has '{collaborator.permission}', needs '{required_permission}'")
                flash(f'❌ You need {required_permission} permission for this', 'danger')
                return redirect(url_for('dashboard'))

            g.user_permission = collaborator.permission
            g.can_edit = collaborator.can_edit()
            g.can_comment = collaborator.can_comment()
            g.can_delete = collaborator.can_delete()
            g.can_manage = collaborator.can_manage_collaborators()

            return view_function(*args, **kwargs)

        return decorated_function

    return decorator


# ----------------- User identification helpers -----------------
def get_user_id():
    """Return a unique id for rate limiting: user_<id> if logged in, otherwise IP."""
    try:
        if getattr(current_user, 'is_authenticated', False):
            return f"user_{current_user.id}"
    except Exception:
        pass
    return get_remote_address()


def get_user_tier():
    """Return user tier for rate limiting decisions: 'premium', 'authenticated', 'anonymous'."""
    try:
        if getattr(current_user, 'is_authenticated', False):
            if getattr(current_user, 'is_premium_active', False) and current_user.is_premium_active():
                return "premium"
            return "authenticated"
    except Exception:
        pass
    return "anonymous"

# Improved limiter configuration
limiter = Limiter(
    key_func=get_user_id,  # IP ke bajaye user-specific identifier use kare
    app=app,
    storage_uri="redis://localhost:6379" if REDIS_AVAILABLE else "memory://",
    default_limits=["1000 per day", "200 per hour"],  # More generous default limits
    headers_enabled=True,  # Response mein rate limit info dikhaye
)

# ------------------- BLUEPRINT REGISTRATION ---------------------

if HAS_AI_API:
    # AI Blueprint ko register karein
    # Saare AI endpoints ab /api/ai prefix se access honge
    app.register_blueprint(ai_blueprint)


# Custom rate limit exceeded handler
from flask import jsonify, render_template, request



def require_auth(f):
    """Require authentication"""
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if not current_user.is_authenticated:
            return jsonify({
                'success': False,
                'message': 'Authentication required'
            }), 401
        return f(*args, **kwargs)
    return decorated_function

def check_rate_limit(f):
    """Check rate limit"""
    @wraps(f)
    def decorated_function(*args, **kwargs):
        user_id = str(current_user.id) if current_user.is_authenticated else 'anonymous'
        allowed, message = rate_limiter.check_limit(
            user_id,
            AIConfig.MAX_REQUESTS_PER_HOUR,
            AIConfig.MAX_REQUESTS_PER_DAY
        )
        if not allowed:
            return jsonify({
                'success': False,
                'message': message
            }), 429
        return f(*args, **kwargs)
    return decorated_function
    

@app.errorhandler(429)
def ratelimit_handler(e):
    logger.warning(f"Rate limit exceeded for {get_remote_address()}: {e.description}")
    
    # Default retry_after value
    retry_after = getattr(e, 'retry_after', None)
    if retry_after is None:
        # Try getting it from headers as fallback
        retry_after = request.headers.get("Retry-After", 60)

    # AJAX or JSON Accept
    if request.headers.get('X-Requested-With') == 'XMLHttpRequest' or \
       'application/json' in request.headers.get('Accept', ''):
        return jsonify({
            "error": "Too many requests",
            "message": "Please try again later.",
            "retry_after": f"{retry_after} seconds",
        }), 429

    # Normal request → Show modal popup via HTML
    return render_template('429.html', retry_after=retry_after), 429


# IP Whitelist filter (apply only in production; configurable)
@limiter.request_filter
def ip_whitelist():
    """
    Trusted IPs ko sirf PRODUCTION me rate limiting se exempt kare.
    Dev/Test me exemption disable rahe taki local par limits test ho sake.
    """
    try:
        env = (current_app.config.get('ENV') or 'production').lower()
    except Exception:
        env = 'production'

    if env != 'production':
        # Dev/Test: NO exemption -> limiter active for localhost too
        return False

    # Production: read allowlist from config (empty means no exemption)
    wl = current_app.config.get('RATE_LIMIT_IP_WHITELIST', [])
    if not isinstance(wl, (list, set, tuple)):
        wl = []
    client_ip = get_remote_address()
    return client_ip in set(wl)

# Admin exemption filter (do NOT exempt security endpoints)
@limiter.request_filter
def admin_exemption():
    try:
        if not current_user.is_authenticated:
            return False
        admin_emails = app.config.get('ADMIN_EMAILS', [])
        if current_user.email not in admin_emails:
            return False
        enforce_on = {'login', 'register', 'verify_note_password',
                      'forgot_password', 'reset_password'}
        ep = (request.endpoint or '').split('.')[-1]
        if ep in enforce_on:
            return False
        return True
    except Exception:
        return False

# ==================== EMAIL RATE LIMITING ====================

# Email rate limiting — uses Redis when available, falls back to in-process dict
_email_send_log: dict = {}   # fallback: {email: [timestamps]}
MAX_EMAILS_PER_USER_PER_HOUR = 100

def check_email_rate_limit(email: str):
    """
    Returns (allowed: bool, remaining: int).
    Uses Redis sliding-window when available, in-process dict otherwise.
    """
    now = datetime.utcnow()

    # ── Redis path ──────────────────────────────────────────
    if REDIS_AVAILABLE and redis_client:
        key = f"email_rl:{email}"
        try:
            pipe = redis_client.pipeline()
            pipe.zadd(key, {str(now.timestamp()): now.timestamp()})
            pipe.zremrangebyscore(key, 0, (now - timedelta(hours=1)).timestamp())
            pipe.zcard(key)
            pipe.expire(key, 3600)
            _, _, count, _ = pipe.execute()
            remaining = max(0, MAX_EMAILS_PER_USER_PER_HOUR - count)
            if count > MAX_EMAILS_PER_USER_PER_HOUR:
                logger.warning(f"Email rate-limit (Redis): {email} — {count} sends/hr")
                return False, 0
            return True, remaining
        except Exception as e:
            logger.warning(f"Redis email-rate check failed, falling back: {e}")

    # ── In-process fallback ──────────────────────────────────
    hour_ago = now - timedelta(hours=1)
    if email not in _email_send_log:
        _email_send_log[email] = []
    recent = [t for t in _email_send_log[email] if t > hour_ago]
    _email_send_log[email] = recent
    if len(recent) >= MAX_EMAILS_PER_USER_PER_HOUR:
        logger.warning(f"Email rate-limit (mem): {email} — {len(recent)} sends/hr")
        return False, 0
    _email_send_log[email].append(now)
    remaining = MAX_EMAILS_PER_USER_PER_HOUR - len(_email_send_log[email])
    return True, remaining


# ==================== EMAIL VALIDATION ====================

def is_valid_email(email):
    """
    Validate email format
    
    Checks:
    - Has @ symbol
    - Has domain name
    - Has TLD (e.g., .com, .org)
    - No spaces
    """
    # Pattern: something@domain.extension
    pattern = r'^[^\s@]+@[^\s@]+\.[^\s@]+$'
    
    is_valid = re.match(pattern, email) is not None
    
    if not is_valid:
        logger.warning(f"❌ Invalid email format: {email}")
    else:
        logger.info(f"✅ Email format valid: {email}")
    
    return is_valid


# ==================== SEND EMAIL HELPER ====================

def send_email_helper(to_email, subject, html_body):
    """
    Send email using Flask-Mail
    
    Args:
        to_email (str): Recipient email address
        subject (str): Email subject
        html_body (str): Email HTML content
    
    Returns:
        (bool, str): (success, error_message)
    """
    try:
        logger.info(f"🔄 Preparing email for: {to_email}")
        logger.info(f"   Subject: {subject[:50]}...")
        
        msg = Message(
            subject=subject,
            recipients=[to_email],
            html=html_body
        )
        
        logger.info(f"📤 Sending email...")
# Send email async — never blocks worker
        _dispatch_email(msg)
        logger.info(f"✅ Note password reset email queued for: {user.email}")
        return True, None
        
    except Exception as e:
        error_msg = str(e)
        logger.error(f"❌ Error sending email to {to_email}: {error_msg}")
        
        # Log additional context
        if "BadCredentials" in error_msg:
            logger.error("   → Issue: Gmail credentials invalid")
            logger.error("   → Solution: Check app password in config.py")
        elif "Connection refused" in error_msg:
            logger.error("   → Issue: SMTP server connection failed")
            logger.error("   → Solution: Check MAIL_SERVER and MAIL_PORT")
        elif "timeout" in error_msg.lower():
            logger.error("   → Issue: Email service timeout")
            logger.error("   → Solution: Increase MAIL_TIMEOUT in config")
        
        return False, error_msg


# ===================================================================
# CSRF helpers (must be defined before any route that uses them)
# ===================================================================

def validate_ajax_csrf():
    """
    Validate CSRF token for AJAX/JSON endpoints that use @csrf.exempt.
    Frontend must send X-CSRFToken header (Flask-WTF token).
    Returns True if valid, False otherwise.
    """
    from flask_wtf.csrf import validate_csrf
    token = (
        request.headers.get('X-CSRFToken') or
        request.headers.get('X-Csrf-Token') or
        (request.get_json(silent=True) or {}).get('csrf_token')
    )
    if not token:
        return False
    try:
        validate_csrf(token)
        return True
    except Exception:
        return False


def require_ajax_csrf(f):
    """Decorator: enforce CSRF on csrf.exempt AJAX routes."""
    @wraps(f)
    def decorated(*args, **kwargs):
        if request.method in ('POST', 'PUT', 'PATCH', 'DELETE'):
            if not validate_ajax_csrf():
                logger.warning(f"CSRF validation failed on {request.endpoint} from {get_remote_address()}")
                return jsonify({'success': False, 'message': 'CSRF token missing or invalid.'}), 403
        return f(*args, **kwargs)
    return decorated


# ===================================================================
# 🔥 API ENDPOINT: Send Invitation Email (FIXED)
# ===================================================================

@app.route('/api/send-invitation-email', methods=['POST'])
@limiter.limit("50 per hour")
@csrf.exempt
@require_ajax_csrf
def send_invitation_email():
    """
    Send invitation email to collaborator
    
    Request:
    {
        "to": "user@example.com",
        "subject": "You've been invited to collaborate!",
        "body": "HTML content",
        "permission": "edit",
        "invitedBy": "John Doe",
        "documentTitle": "Project Report"
    }
    
    Response:
    {
        "success": true,
        "message": "Email sent successfully",
        "messageId": "uuid",
        "provider": "flask-mail",
        "timestamp": datetime.utcnow().isoformat() + "Z"
    }
    """
    
    try:
        logger.debug("\n" + "="*60)
        logger.debug("📧 SEND INVITATION EMAIL REQUEST")
        logger.debug("="*60)
        
        data = request.get_json()
        
        if not data:
            logger.warning("No JSON data provided")
            return jsonify({
                'success': False,
                'error': 'No data provided'
            }), 400
        
        # Validate required fields
        required = ['to', 'subject', 'body']
        missing = [f for f in required if not data.get(f)]
        
        if missing:
            logger.warning(f"Missing required fields: {missing}")
            return jsonify({
                'success': False,
                'error': f'Missing: {", ".join(missing)}'
            }), 400
        
        to_email = data.get('to', '').strip().lower()
        subject = data.get('subject', '').strip()
        html_body = data.get('body', '').strip()
        permission = data.get('permission', 'edit')
        invited_by = data.get('invitedBy', 'Unknown')
        document_title = data.get('documentTitle', 'Document')
        
        logger.debug(f"To: {to_email}")
        logger.debug(f"Subject: {subject[:60]}...")
        logger.debug(f"Permission: {permission}")
        logger.debug(f"Invited by: {invited_by}")
        logger.debug("="*60)
        
        # ===== STEP 1: Validate email format =====
        logger.debug("1️⃣ Validating email format...")
        if not is_valid_email(to_email):
            logger.error(f"Invalid email format: {to_email}")
            return jsonify({
                'success': False,
                'error': 'Invalid email address format'
            }), 400
        logger.info("   ✅ Email format valid")
        
        # ===== STEP 2: Check rate limit =====
        logger.debug("2️⃣ Checking rate limit...")
        allowed, remaining = check_email_rate_limit(to_email)
        
        if not allowed:
            logger.warning(f"Rate limit exceeded for {to_email}")
            return jsonify({
                'success': False,
                'error': 'Too many invitations sent to this email. Try again later.',
                'reset_time': (datetime.utcnow() + timedelta(hours=1)).isoformat(),
                'remaining': 0
            }), 429
        logger.info(f"   ✅ Rate limit OK (Remaining: {remaining})")
        
        # ===== STEP 3: Validate HTML body length =====
        logger.debug("3️⃣ Validating email body...")
        if len(html_body) < 10:
            logger.warning("Email body too short")
            return jsonify({
                'success': False,
                'error': 'Email body content is too short'
            }), 400
        logger.info(f"   ✅ Body length OK ({len(html_body)} chars)")
        
        # ===== STEP 4: Send the email =====
        logger.debug("4️⃣ Sending email...")
        success, error = send_email_helper(to_email, subject, html_body)
        
        if success:
            logger.info("✅ EMAIL SENT SUCCESSFULLY")
            logger.debug("="*60 + "\n")
            
            # Persist a pending collaborator/invitation if note_id supplied
            try:
                note_id = int(data.get('note_id')) if data.get('note_id') else None
            except Exception:
                note_id = None

            if note_id:
                try:
                    # Avoid duplicate pending invitations for same email+note
                    existing = Collaborator.query.filter_by(
                        note_id=note_id,
                        collaborator_email=to_email
                    ).first()

                    if existing:
                        # Update permission/status if needed
                        existing.permission = permission
                        existing.status = existing.status or 'pending'
                        db.session.add(existing)
                    else:
                        collab = Collaborator(
                            note_id=note_id,
                            collaborator_email=to_email,
                            permission=permission,
                            status='pending'
                        )
                        db.session.add(collab)

                    db.session.commit()
                    logger.info(f"Saved pending collaborator for {to_email} on note {note_id}")
                except Exception as e:
                    db.session.rollback()
                    logger.error(f"Failed to save pending collaborator: {e}")

            return jsonify({
                'success': True,
                'message': 'Email sent successfully',
                'messageId': f"inv_{int(datetime.utcnow().timestamp())}",
                'provider': 'flask-mail',
                'timestamp': datetime.utcnow().isoformat(),
                'remaining': remaining
            }), 200
        else:
            logger.error(f"❌ EMAIL SEND FAILED: {error}")
            logger.debug("="*60 + "\n")
            
            logger.error(f"❌ Email send failed: {error}")
            return jsonify({
                'success': False,
                'error': error or 'Failed to send email',
                'details': error
            }), 500
            
    except Exception as e:
        logger.error(f"❌ EXCEPTION: {type(e).__name__}: {str(e)}")
        logger.debug("="*60 + "\n")
        
        logger.error(f"❌ Exception in send_invitation_email: {str(e)}")
        logger.error(f"Traceback: {traceback.format_exc()}")
        
        return jsonify({
            'success': False,
            'error': 'Internal server error',
            'details': str(e) if app.debug else None
        }), 500


# ===================================================================
# 📧 API ENDPOINT: Email Health Check
# ===================================================================

@app.route('/api/email-health', methods=['GET'])
@limiter.exempt
def email_health():
    """Check if email service is configured"""
    
    try:
        mail_server = app.config.get('MAIL_SERVER')
        mail_port = app.config.get('MAIL_PORT')
        mail_username = app.config.get('MAIL_USERNAME')
        mail_use_tls = app.config.get('MAIL_USE_TLS', False)
        
        mail_configured = bool(mail_server and mail_port and mail_username)
        
        status = "✅ Configured" if mail_configured else "❌ Not Configured"
        
        logger.info(f"📊 Email Health Check: {status}")
        logger.info(f"   Server: {mail_server or 'Not set'}")
        logger.info(f"   Port: {mail_port or 'Not set'}")
        logger.info(f"   Username: {mail_username[:10] + '***' if mail_username else 'Not set'}")
        logger.info(f"   TLS: {mail_use_tls}")
        
        return jsonify({
            'success': True,
            'mail_configured': mail_configured,
            'mail_server': mail_server or 'Not configured',
            'mail_port': mail_port or 'Not set',
            'mail_use_tls': mail_use_tls,
            'mail_username_set': bool(mail_username),
            'status': status,
            'timestamp': datetime.utcnow().isoformat()
        }), 200
        
    except Exception as e:
        logger.error(f"Error checking email health: {e}")
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500


# ===================================================================
# 📧 API ENDPOINT: Email Status (requires authentication)
# ===================================================================

@app.route('/api/email-status', methods=['GET'])
@login_required
@limiter.limit("10 per minute")
@csrf.exempt
@require_ajax_csrf
def email_status():
    """Check current user's email usage"""
    
    try:
        user_email = current_user.email
        allowed, remaining = check_email_rate_limit(user_email)
        
        logger.info(f"📊 Email status for {user_email}: Allowed={allowed}, Remaining={remaining}")
        
        return jsonify({
            'success': True,
            'email': user_email,
            'can_send': allowed,
            'remaining': remaining,
            'max_per_hour': MAX_EMAILS_PER_USER_PER_HOUR,
            'reset_time': (datetime.utcnow() + timedelta(hours=1)).isoformat()
        }), 200
        
    except Exception as e:
        logger.error(f"❌ Error in email_status: {str(e)}")
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500


# ===================================================================
# 📧 API ENDPOINT: Get Pending Invitations
# ===================================================================

@app.route('/api/pending-invitations', methods=['GET'])
@login_required
@limiter.limit("10 per minute")
@csrf.exempt
@require_ajax_csrf
def pending_invitations():
    """Get pending collaboration invitations for current user"""
    
    try:
        # Query real pending Collaborator records where this user is invited
        pending_collabs = Collaborator.query.filter_by(
            collaborator_email=current_user.email,
            status='pending'
        ).all()

        pending = []
        for c in pending_collabs:
            if c.is_expired():
                continue
            note = Note.query.get(c.note_id)
            inviter = User.query.get(note.user_id) if note else None
            pending.append({
                'id':             c.id,
                'note_id':        c.note_id,
                'note_title':     note.title if note else 'Untitled',
                'invited_by':     inviter.username if inviter else 'Unknown',
                'permission':     c.permission,
                'invited_at':     c.invited_at.isoformat() if hasattr(c, 'invited_at') and c.invited_at else None,
                'accept_url':     url_for('accept_invitation', email=current_user.email, _external=True),
            })

        logger.info(f"Pending invitations for {current_user.email}: {len(pending)}")

        return jsonify({
            'success': True,
            'pending_invitations': pending,
            'count': len(pending)
        }), 200
        
    except Exception as e:
        logger.error(f"❌ Error in pending_invitations: {str(e)}")
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500


# ===================================================================
# 📚 API ENDPOINT: Collaboration API Documentation
# ===================================================================

@app.route('/api/collaboration-docs', methods=['GET'])
@login_required
def collaboration_docs():
    """API documentation"""
    
    docs = {
        'title': 'Collaboration System API (REAL MODE ONLY)',
        'version': '1.0.0',
        'endpoints': {
            'send_email': {
                'url': '/api/send-invitation-email',
                'method': 'POST',
                'rate_limit': '50 per hour',
                'description': 'Send invitation email to collaborator',
                'required_fields': ['to', 'subject', 'body']
            },
            'email_health': {
                'url': '/api/email-health',
                'method': 'GET',
                'rate_limit': 'None',
                'description': 'Check if email service is configured'
            },
            'email_status': {
                'url': '/api/email-status',
                'method': 'GET',
                'requires_auth': True,
                'description': 'Check current user email rate limit status'
            },
            'pending_invitations': {
                'url': '/api/pending-invitations',
                'method': 'GET',
                'requires_auth': True,
                'description': 'Get pending invitations for current user'
            }
        },
        'error_codes': {
            400: 'Bad Request - Missing or invalid parameters',
            401: 'Unauthorized - Authentication required',
            429: 'Too Many Requests - Rate limit exceeded',
            500: 'Internal Server Error'
        }
    }
    
    return jsonify(docs), 200
# ===================================================================
# ADD THESE ROUTES TO YOUR app.py
# Place them after the /api/send-invitation-email route (around line 330)
# ===================================================================

# ===================================================================
# 🎯 PUBLIC ROUTE: Accept Invitation (No authentication required)
# ===================================================================

@app.route('/accept-invitation', methods=['GET', 'POST'])
@limiter.limit("20 per hour")  # Prevent abuse
def accept_invitation():
    """
    Accept a collaboration invitation.
    Called from email link or manually by user.
    """
    try:
        email = request.args.get('email', '').strip()
        token = request.args.get('token', '').strip()
        
        # If POST request, get data from form
        if request.method == 'POST':
            data = request.get_json() or request.form
            email = data.get('email', '').strip()
            token = data.get('token', '').strip()

        if not email:
            logger.warning("Accept invitation attempted without email")
            if request.is_json:
                return jsonify({
                    'success': False,
                    'message': 'Email address required'
                }), 400
            flash('❌ Email address is required', 'danger')
            return redirect(url_for('login'))

        # Check if user is logged in
        if not current_user.is_authenticated:
            # Store email in session for after login
            session['pending_invitation_email'] = email
            flash(f'📧 Please log in to accept the invitation to {email}', 'info')
            return redirect(url_for('login'))

        # Log the acceptance
        logger.info(f"Invitation accepted: {email} by user {current_user.username}")

        # ✅ IMPLEMENTATION: Find Collaborator pending record and mark accepted
        try:
            collab = Collaborator.query.filter_by(
                collaborator_email=email,
                status='pending'
            ).order_by(Collaborator.invited_at.desc()).first()

            if collab:
                if collab.is_expired():
                    logger.warning(f"Invitation expired for {email} on note {collab.note_id}")
                    flash('❌ This invitation has expired', 'danger')
                    return redirect(url_for('dashboard'))

                # Ensure note exists
                if not collab.note:
                    logger.error(f"Collaborator record has no note: {collab}")
                # Link to current_user if registered and mark accepted
                collab.collaborator_id = current_user.id
                collab.status = 'accepted'
                collab.accepted_at = datetime.utcnow()
                db.session.add(collab)
                db.session.commit()
                logger.info(f"✅ Collaborator record accepted: {email} for note {collab.note_id}")
            else:
                logger.info(f"No pending collaborator record found for {email}")
        except Exception as e:
            db.session.rollback()
            logger.error(f"Error marking collaborator accepted: {e}")

        if request.is_json:
            return jsonify({
                'success': True,
                'message': f'✅ Invitation accepted! Welcome to collaboration.',
                'redirect_url': url_for('dashboard')
            }), 200

        flash(f'✅ Invitation accepted! You are now a collaborator.', 'success')
        return redirect(url_for('dashboard'))

    except Exception as e:
        logger.error(f"Error in accept_invitation: {e}")
        if request.is_json:
            return jsonify({
                'success': False,
                'message': 'Error accepting invitation'
            }), 500
        flash('❌ Error accepting invitation', 'danger')
        return redirect(url_for('login'))


# ===================================================================
# 📧 ROUTE: Get Pending Invitations for Current User
# ===================================================================

@app.route('/api/user-pending-invitations')
@login_required
@limiter.limit("20 per minute")
def user_pending_invitations():
    """Get all pending invitations for the logged-in user."""
    try:
        pending_collabs = Collaborator.query.filter_by(
            collaborator_email=current_user.email,
            status='pending'
        ).all()

        invitations_data = []
        for c in pending_collabs:
            if c.is_expired():
                continue
            note = Note.query.get(c.note_id)
            inviter = User.query.get(note.user_id) if note else None
            invitations_data.append({
                'id':         c.id,
                'note_id':    c.note_id,
                'from':       inviter.username if inviter else 'Unknown',
                'document':   note.title if note else 'Untitled',
                'permission': c.permission,
                'sent_at':    c.invited_at.isoformat() if hasattr(c, 'invited_at') and c.invited_at else None,
            })

        return jsonify({
            'success': True,
            'pending_invitations': invitations_data,
            'count': len(invitations_data)
        }), 200

    except Exception as e:
        logger.error(f"Error fetching pending invitations: {e}")
        return jsonify({'success': False, 'message': 'Error loading invitations'}), 500


# ===================================================================
# 🚫 ROUTE: Reject/Decline Invitation
# ===================================================================

@app.route('/api/decline-invitation', methods=['POST'])
@login_required
@limiter.limit("10 per minute")
def decline_invitation():
    """Decline a collaboration invitation."""
    try:
        data = request.get_json() or {}
        invitation_id = data.get('invitation_id')

        if not invitation_id:
            return jsonify({'success': False, 'message': 'Invitation ID required'}), 400

        collab = Collaborator.query.filter_by(
            id=invitation_id,
            collaborator_email=current_user.email,
            status='pending'
        ).first()

        if not collab:
            return jsonify({'success': False, 'message': 'Invitation not found or already processed'}), 404

        collab.status = 'declined'
        db.session.commit()

        logger.info(f"Invitation {invitation_id} declined by {current_user.username}")
        return jsonify({'success': True, 'message': 'Invitation declined'}), 200

    except Exception as e:
        db.session.rollback()
        logger.error(f"Error declining invitation: {e}")
        return jsonify({'success': False, 'message': 'Error declining invitation'}), 500


# ===================================================================
# 📧 ROUTE: Resend Invitation Email
# ===================================================================

@app.route('/api/resend-invitation', methods=['POST'])
@login_required
@limiter.limit("5 per hour")
def resend_invitation():
    """Resend an invitation email to a collaborator"""
    try:
        data = request.get_json()
        email = data.get('email', '').strip()

        if not email or not is_valid_email(email):
            return jsonify({
                'success': False,
                'message': 'Valid email required'
            }), 400

        # Check rate limit for email
        allowed, remaining = check_email_rate_limit(email)
        if not allowed:
            return jsonify({
                'success': False,
                'message': 'Too many invitations. Try again later.',
                'reset_time': (datetime.utcnow() + timedelta(hours=1)).isoformat()
            }), 429

        # Verify the collaborator record still exists and is pending
        collab = Collaborator.query.filter_by(
            collaborator_email=email,
            status='pending'
        ).filter(
            Collaborator.note_id.in_(
                db.session.query(Note.id).filter_by(user_id=current_user.id)
            )
        ).first()

        if not collab:
            return jsonify({'success': False, 'message': 'No pending invitation found for this email'}), 404

        note = Note.query.get(collab.note_id)
        subject = f"Reminder: Invitation to collaborate on '{note.title if note else 'a note'}'"
        html_body = f"""
        <html>
        <body style="font-family: Arial, sans-serif;">
            <div style="max-width: 600px; margin: 0 auto; padding: 20px;">
                <h2>📧 Collaboration Invitation Reminder</h2>
                <p>Hi,</p>
                <p>{current_user.username} reminded you about your pending collaboration invitation.</p>
                <div style="margin: 30px 0; text-align: center;">
                    <a href="{url_for('accept_invitation', email=email, _external=True)}"
                       style="background: #28a745; color: white; padding: 12px 25px; 
                              text-decoration: none; border-radius: 5px; display: inline-block;">
                        ✅ Accept Invitation
                    </a>
                </div>
                <p style="color: #666; font-size: 12px;">
                    Invitation expires in 7 days.
                </p>
            </div>
        </body>
        </html>
        """

        success, error = send_email_helper(email, subject, html_body)

        if success:
            logger.info(f"Invitation resent to {email} by user {current_user.username}")
            return jsonify({
                'success': True,
                'message': f'✅ Invitation reminder sent to {email}'
            }), 200
        else:
            logger.error(f"Failed to resend invitation: {error}")
            return jsonify({
                'success': False,
                'message': f'Failed to send email: {error}'
            }), 500

    except Exception as e:
        logger.error(f"Error in resend_invitation: {e}")
        return jsonify({
            'success': False,
            'message': 'Error resending invitation'
        }), 500


# ===================================================================
# 🔍 ROUTE: Get Invitation Details (Public)
# ===================================================================

@app.route('/api/invitation-details/<token>')
@limiter.limit("10 per minute")
def invitation_details(token):
    """Decode invitation token and return real details."""
    try:
        try:
            payload = serializer.loads(token, salt='invitation-salt', max_age=604800)  # 7 days
        except SignatureExpired:
            return jsonify({'success': False, 'message': 'Invitation link has expired'}), 410
        except BadSignature:
            return jsonify({'success': False, 'message': 'Invalid invitation link'}), 400

        collab_id = payload.get('collab_id')
        collab = Collaborator.query.get(collab_id) if collab_id else None

        if not collab or collab.status != 'pending':
            return jsonify({'success': False, 'message': 'Invitation not found or already used'}), 404

        note  = Note.query.get(collab.note_id)
        inviter = User.query.get(note.user_id) if note else None

        return jsonify({
            'success': True,
            'invitation': {
                'id':         collab.id,
                'from':       inviter.username if inviter else 'Unknown',
                'document':   note.title if note else 'Untitled',
                'permission': collab.permission,
                'email':      collab.collaborator_email,
                'expired':    collab.is_expired(),
            }
        }), 200

    except Exception as e:
        logger.error(f"Error fetching invitation details: {e}")
        return jsonify({'success': False, 'message': 'Invalid or expired invitation'}), 404


# ===================================================================
# 🎯 ROUTE: Accept Invitation After Login (Automatic)
# ===================================================================

@app.before_request
def handle_pending_invitation():
    """Check for pending invitation after user logs in"""
    if current_user.is_authenticated:
        pending_email = session.get('pending_invitation_email')
        if pending_email:
            session.pop('pending_invitation_email', None)
            logger.info(f"Processing pending invitation for {pending_email}")
            # Optionally, auto-accept or redirect to confirmation page


# ===================================================================
# 🗑️ ROUTE: Remove Collaborator
# ===================================================================

@app.route('/api/remove-collaborator/<int:note_id>', methods=['POST'])
@login_required
@limiter.limit("20 per minute")
@check_collaboration_permission
@require_permission('admin')
def remove_collaborator(note_id):
    """Remove a collaborator from a document (admin only)"""
    try:
        data = request.get_json() or {}
        email = data.get('email', '').strip()

        if not email:
            return jsonify({
                'success': False,
                'message': 'Email required'
            }), 400

        # Verify note exists
        note = Note.query.get(note_id)
        if not note:
            return jsonify({'success': False, 'message': 'Note not found'}), 404

        # Remove collaborator record if exists
        collaborator = Collaborator.query.filter_by(
            note_id=note_id,
            collaborator_email=email
        ).first()

        if collaborator:
            db.session.delete(collaborator)
            db.session.commit()
            logger.info(f"Collaborator {email} removed from note {note_id} by user {current_user.username}")
            return jsonify({
                'success': True,
                'message': f'✅ {email} has been removed as a collaborator'
            }), 200

        return jsonify({'success': False, 'message': 'Collaborator not found'}), 404

    except Exception as e:
        logger.error(f"Error removing collaborator: {e}")
        return jsonify({
            'success': False,
            'message': 'Error removing collaborator'
        }), 500


# ===================================================================
# 🔗 ROUTE: Get Collaborators for a Note
# ===================================================================

@app.route('/api/note/<int:note_id>/collaborators')
@login_required
@limiter.limit("30 per minute")
def get_note_collaborators(note_id):
    """Get real list of collaborators for a specific note."""
    try:
        note = Note.query.get(note_id)
        if not note:
            return jsonify({'success': False, 'message': 'Note not found'}), 404

        if note.user_id != current_user.id:
            return jsonify({'success': False, 'message': 'Unauthorized'}), 403

        collaborators = Collaborator.query.filter_by(note_id=note_id).all()

        collaborators_data = [{
            'email': current_user.email,
            'name': getattr(current_user, 'username', current_user.email),
            'permission': 'owner',
            'status': 'active'
        }] + [{
            'email': c.collaborator_email,
            'name': c.collaborator_email,
            'permission': c.permission,
            'status': c.status
        } for c in collaborators]

        return jsonify({
            'success': True,
            'collaborators': collaborators_data,
            'count': len(collaborators_data)
        }), 200

    except Exception as e:
        logger.error(f"Error fetching collaborators: {e}")
        return jsonify({'success': False, 'message': 'Error loading collaborators'}), 500



# ===================================================================
# 📊 ROUTE: Collaboration Activity Log
# ===================================================================

@app.route('/api/collaboration/activity/<int:note_id>')
@login_required
@limiter.limit("30 per minute")
def collaboration_activity(note_id):
    """Get collaboration activity for a note. Implement ActivityLog model for full history."""
    try:
        note = Note.query.get(note_id)
        if not note:
            return jsonify({'success': False, 'message': 'Note not found'}), 404

        # Basic activity: show note creation + last edit
        activity_data = [
            {
                'timestamp': note.created_at.isoformat(),
                'user': current_user.username,
                'action': 'Created note',
                'details': note.title,
            }
        ]
        if note.updated_at != note.created_at:
            activity_data.append({
                'timestamp': note.updated_at.isoformat(),
                'user': current_user.username,
                'action': 'Last edited',
                'details': note.title,
            })

        return jsonify({
            'success': True,
            'activity': activity_data,
            'count': len(activity_data),
        }), 200

    except Exception as e:
        logger.error(f"Error fetching collaboration activity: {e}")
        return jsonify({'success': False, 'message': 'Error loading activity'}), 500
    
@app.route('/api/send-invitation-with-permission', methods=['POST'])
@login_required
@limiter.limit("50 per hour")
@csrf.exempt
@require_ajax_csrf
def send_invitation_with_permission():
    """
    Send invitation with specific permission
    
    Request:
    {
        "note_id": 1,
        "email": "user@example.com",
        "permission": "admin" | "edit" | "comment" | "view"
    }
    """
    try:
        data = request.get_json()
        note_id = data.get('note_id')
        email = data.get('email', '').strip().lower()
        permission = data.get('permission', 'view').lower()
        
        # Validation
        if not note_id or not email:
            return jsonify({'success': False, 'error': 'Missing data'}), 400
        
        if permission not in ['admin', 'edit', 'comment', 'view']:
            return jsonify({'success': False, 'error': 'Invalid permission'}), 400
        
        # Verify ownership
        note = Note.query.get_or_404(note_id)
        if note.user_id != current_user.id:
            return jsonify({'success': False, 'error': 'Unauthorized'}), 403
        
        # Get current page URL
        current_page = request.referrer or request.path
        
        # Create collaborator record
        collaborator = Collaborator(
            note_id=note_id,
            collaborator_email=email,
            permission=permission,
            restricted_page=current_page if permission == 'view' else None,
            status='pending'
        )
        
        db.session.add(collaborator)
        db.session.commit()
        
        # Send email notification
        subject = f"Invited to collaborate on: {note.title}"
        
        permission_info = {
            'admin': '⚙️ Full access (Admin)',
            'edit': '✏️ Can edit',
            'comment': '💬 Can comment',
            'view': '👁️ View only (This page only)'
        }
        
        body = f"""
        <html>
        <body style="font-family: Arial; color: #333;">
            <h2>👋 You're invited to collaborate!</h2>
            
            <p><strong>{current_user.username}</strong> invited you to collaborate on 
            <strong>"{note.title}"</strong>.</p>
            
            <div style="background: #f0f0f0; padding: 15px; border-left: 4px solid #007bff; margin: 20px 0;">
                <p><strong>🔐 Your Permission: {permission_info.get(permission, permission)}</strong></p>
                
                {'<p style="color: red;"><strong>⚠️ Note:</strong> आप सिर्फ इस page तक access कर सकते हैं।</p>' if permission == 'view' else ''}
            </div>
            
            <p>
                <a href="{url_for('accept_invitation', email=email, _external=True)}" 
                   style="background: #28a745; color: white; padding: 12px 25px; text-decoration: none; border-radius: 5px;">
                    ✅ Accept Invitation
                </a>
            </p>
        </body>
        </html>
        """
        
        success, error = send_email_helper(email, subject, body)
        
        if success:
            logger.info(f"✅ Invitation sent to {email} with permission: {permission}")
            return jsonify({
                'success': True,
                'message': f'✅ Invitation sent! Permission: {permission}'
            }), 200
        else:
            return jsonify({'success': False, 'error': error}), 500
    
    except Exception as e:
        logger.error(f"Error: {e}")
        return jsonify({'success': False, 'error': str(e)}), 500
    
# ==================== ADD THESE FUNCTIONS TO app.py ====================
# Add these after imports, around line 250 (after utility_processor)

def clean_html(raw_html: str) -> str:
    """
    Sanitize HTML content while PRESERVING safe structure.
    Removes dangerous scripts, avoids Bleach warnings,
    and works without tinycss2.
    """

    if not raw_html or not raw_html.strip():
        return '<p><br></p>'

    # ================= ALLOWED TAGS =================
    ALLOWED_TAGS = [
        'p', 'br', 'span', 'div',
        'b', 'strong', 'i', 'em', 'u', 's', 'mark', 'sub', 'sup',
        'small', 'del', 'ins',
        'h1', 'h2', 'h3', 'h4', 'h5', 'h6',
        'ul', 'ol', 'li',
        'a', 'img',
        'table', 'thead', 'tbody', 'tfoot', 'tr', 'th', 'td',
        'blockquote', 'pre', 'code', 'hr'
    ]

    # ======================================================
    # ALLOWED ATTRIBUTES
    # 🔧 HINDI FIX: style, lang, dir preserved for
    #    Devanagari/Indic font-family, text-direction etc.
    # ======================================================
    ALLOWED_ATTRIBUTES = {
        '*': ['class', 'style', 'lang', 'dir'],
        'a': ['href', 'title', 'target', 'rel'],
        'img': ['src', 'alt', 'width', 'height'],
        'td': ['colspan', 'rowspan'],
        'th': ['colspan', 'rowspan', 'scope'],
        'table': ['border', 'cellpadding', 'cellspacing'],
    }

    # ================= PRE-CLEAN DANGEROUS CONTENT =================
    raw_html = re.sub(
        r'<script\b[^<]*(?:(?!<\/script>)<[^<]*)*<\/script>',
        '',
        raw_html,
        flags=re.IGNORECASE
    )

    raw_html = re.sub(
        r'<style\b[^<]*(?:(?!<\/style>)<[^<]*)*<\/style>',
        '',
        raw_html,
        flags=re.IGNORECASE
    )

    raw_html = re.sub(
        r'\son\w+\s*=\s*["\'][^"\']*["\']',
        '',
        raw_html,
        flags=re.IGNORECASE
    )

    # ================= BLEACH CLEAN (SAFE, NO CSS NEEDED) =================
    try:
        cleaner = bleach.Cleaner(
            tags=ALLOWED_TAGS,
            attributes=ALLOWED_ATTRIBUTES,
            strip=True
        )
        cleaned_html = cleaner.clean(raw_html)

    except Exception as e:
        logger.error(f"Bleach failed — fallback to BeautifulSoup: {e}")

        soup = BeautifulSoup(raw_html, 'html.parser')
        for tag in soup.find_all(['script', 'style', 'iframe', 'object', 'embed']):
            tag.decompose()

        cleaned_html = str(soup)

    # ================= FINAL CLEANUP =================
    cleaned_html = re.sub(r'\n{3,}', '\n\n', cleaned_html).strip()

    if not cleaned_html or cleaned_html.strip() in ['', '<p></p>', '<div></div>']:
        return '<p><br></p>'

    return cleaned_html



def html_to_plain_lines(html_content):
    """
    Convert HTML to plain text lines while preserving structure.
    Returns list of tuples: (text, alignment, is_bold, is_italic)
    """
    from bs4 import BeautifulSoup, NavigableString
    
    soup = BeautifulSoup(html_content, 'html.parser')
    
    # Remove unwanted tags
    for tag in soup.find_all(['script', 'style', 'iframe', 'object', 'embed']):
        tag.decompose()
    
    lines = []
    
    def extract_lines(node, parent_format=None):
        """Recursively extract text lines with formatting"""
        if parent_format is None:
            parent_format = {'bold': False, 'italic': False, 'align': 'left'}
        
        # Text node
        if isinstance(node, NavigableString):
            text = str(node).strip()
            if text:
                lines.append({
                    'text': text,
                    'bold': parent_format['bold'],
                    'italic': parent_format['italic'],
                    'align': parent_format['align']
                })
            return
        
        # Element node
        if node.name is None:
            return
        
        # Block elements - create new line
        if node.name in ['p', 'div', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
            # Get alignment from style
            style = node.get('style', '')
            align = 'left'
            if 'text-align' in style:
                if 'right' in style:
                    align = 'right'
                elif 'center' in style:
                    align = 'center'
                elif 'justify' in style:
                    align = 'justify'
            
            # Update format
            new_format = parent_format.copy()
            new_format['align'] = align
            
            # Extract text
            text = node.get_text(separator=' ', strip=True)
            if text:
                lines.append({
                    'text': text,
                    'bold': node.name.startswith('h') or parent_format['bold'],
                    'italic': parent_format['italic'],
                    'align': align,
                    'is_heading': node.name.startswith('h')
                })
            return
        
        # Formatting elements
        elif node.name in ['strong', 'b']:
            new_format = parent_format.copy()
            new_format['bold'] = True
            for child in node.children:
                extract_lines(child, new_format)
            return
        
        elif node.name in ['em', 'i']:
            new_format = parent_format.copy()
            new_format['italic'] = True
            for child in node.children:
                extract_lines(child, new_format)
            return
        
        # Line break
        elif node.name == 'br':
            lines.append({'text': '', 'bold': False, 'italic': False, 'align': 'left'})
            return
        
        # Lists
        elif node.name in ['ul', 'ol']:
            for idx, li in enumerate(node.find_all('li', recursive=False), 1):
                text = li.get_text(strip=True)
                if text:
                    prefix = '• ' if node.name == 'ul' else f'{idx}. '
                    lines.append({
                        'text': prefix + text,
                        'bold': False,
                        'italic': False,
                        'align': 'left'
                    })
            return
        
        # Process children
        for child in node.children:
            extract_lines(child, parent_format)
    
    # Extract from body
    body = soup.find('body') or soup
    for child in body.children:
        extract_lines(child)
    
    return lines


def get_note_preview(html_content, max_length=150):
    """
    Convert HTML content to clean preview text for dashboard
    """
    if not html_content:
        return "Empty note"
    
    try:
        # Parse HTML
        soup = BeautifulSoup(html_content, 'html.parser')
        
        # Remove script and style tags
        for tag in soup(['script', 'style', 'meta', 'link']):
            tag.decompose()
        
        # Get plain text
        text = soup.get_text(separator=' ', strip=True)
        
        # Clean up whitespace
        text = re.sub(r'\s+', ' ', text).strip()
        
        # Remove date/time patterns
        date_time_pattern = r"\b\d{1,2} (?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]* \d{4},? \d{1,2}:\d{2}(?::\d{2})? ?[APMapm]{2}\b"
        text = re.sub(date_time_pattern, '', text)
        
        date_only_pattern = r"\b\d{1,2} (?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]* \d{4}\b"
        text = re.sub(date_only_pattern, '', text)
        
        # Clean up again
        text = re.sub(r'\s+', ' ', text).strip()
        
        # Truncate to max length
        if len(text) > max_length:
            text = text[:max_length].rsplit(' ', 1)[0] + '...'
        
        return text if text else "Empty note"
        
    except Exception as e:
        logger.error(f"Error in get_note_preview: {e}")
        return "Preview unavailable"


# Update your context processor (around line 100)
@app.context_processor
def utility_processor():
    """Make utility functions available in templates"""
    return dict(
        get_note_preview=get_note_preview,
        csrf_token=generate_csrf(),
        timedelta=timedelta   # IST conversion ke liye templates mein
    )


def premium_required(f):
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if not current_user.is_authenticated or not current_user.is_premium_active():
            flash('This feature requires Premium membership!', 'warning')
            return redirect(url_for('premium'))
        return f(*args, **kwargs)
    return decorated_function




def send_reset_email(user):
    token = serializer.dumps(user.email, salt='password-reset-salt')
    reset_url = url_for('reset_password', token=token, _external=True)

    msg = Message(
        "Password Reset Request",
        sender=app.config['MAIL_DEFAULT_SENDER'],
        recipients=[user.email]
    )
    msg.body = (
        f"To reset your password, visit:\n\n{reset_url}\n\n"
        "If you did not request this, ignore this email."
    )
    thread = threading.Thread(target=_send_mail_async, args=(app, msg))
    thread.daemon = True
    thread.start()
    flash('A password reset link has been sent to your email.', 'info')

# Helper function for note access control
def check_note_access(note_id):
    """
    Check if current user has access to a note.
    Supports both owner and accepted collaborators.
    Returns: (note, has_access)
    """
    # Try owner first
    note = Note.query.filter_by(id=note_id, user_id=current_user.id).first()

    if note:
        # Owner — check private password if set
        if not note.is_private:
            return note, True
        if session.get(f'note_verified_{note_id}', False):
            return note, True
        return note, False

    # Try collaborator access
    collab = Collaborator.query.filter_by(
        note_id=note_id,
        collaborator_id=current_user.id,
        status='accepted'
    ).first()

    if collab and not collab.is_expired():
        note = Note.query.get(note_id)
        return note, True

    return None, False

def apply_tiered_rate_limit(route_func, anonymous_limit, auth_limit, premium_limit):
    """Apply different rate limits based on user tier"""
    user_tier = get_user_tier()
    
    if user_tier == "premium":
        return limiter.limit(premium_limit)(route_func)
    elif user_tier == "authenticated":
        return limiter.limit(auth_limit)(route_func)
    else:
        return limiter.limit(anonymous_limit)(route_func)

# ------------------- Routes ---------------------




@app.route('/')
@limiter.limit("100 per minute")  # General browsing limit
def index():
    if current_user.is_authenticated:
        return redirect(url_for('dashboard'))
    return render_template('index.html')
    
@app.route('/note')
@login_required
def note_page():
    return render_template("note.html")
    
# Version control को database में save करने के लिए
@app.route('/api/version/save', methods=['POST'])
@login_required
def save_version():
    """Placeholder — implement with Version model when ready."""
    return jsonify({'success': False, 'message': 'Not implemented yet.'}), 501

# Comments को persist करने के लिए
@app.route('/api/comments/save', methods=['POST'])
@login_required
def save_comments():
    """Placeholder — implement with Comment model when ready."""
    return jsonify({'success': False, 'message': 'Not implemented yet.'}), 501


# Real collaboration के लिए WebSocket
@socketio.on('content_change')
def handle_content_change(data):
    # Broadcast to all connected users
    emit('content_update', data, broadcast=True)


def generate_otp():
    """Generates a 6-digit OTP."""
    return str(randint(100000, 999999))

def is_otp_valid(email, otp_code):
    """
    Checks if the provided OTP is valid and not expired.
    Uses timezone-naive comparison (both DB and now() are naive UTC).
    """
    user = User.query.filter_by(email=email).first()
    if not user:
        return False
    if not user.otp_code or not user.otp_expiry:
        return False
    # Both values are naive UTC — compare directly
    if user.otp_code == str(otp_code) and user.otp_expiry > datetime.utcnow():
        return True
    return False

def _send_mail_async(flask_app, msg):
    """Send a Flask-Mail message in a background thread."""
    with flask_app.app_context():
        try:
            mail.send(msg)
            logger.info(f"Async email sent to {msg.recipients}")
        except Exception as e:
            logger.error(f"Async email failed to {msg.recipients}: {e}")


def send_verification_email(user, otp_code):
    """Queue OTP email in background thread — never blocks the web worker."""
    try:
        msg = Message(
            "NoteSaver Pro: Email Verification Code",
            sender=app.config['MAIL_DEFAULT_SENDER'],
            recipients=[user.email]
        )
        msg.body = (
            f"Hello {user.username},\n\n"
            f"Your verification code is: {otp_code}\n\n"
            "This code expires in 5 minutes. "
            "If you did not request this, ignore this email.\n\n"
            "Thank you,\nNoteSaver Pro Team"
        )
        thread = threading.Thread(target=_send_mail_async, args=(app, msg))
        thread.daemon = True
        thread.start()
        return True
    except Exception as e:
        logger.error(f"Error queuing verification email to {user.email}: {e}")
        return False
        
def send_username_reminder_email(user):
    try:
        msg = Message(
            "NoteSaver Pro: Your Username Reminder",
            sender=app.config['MAIL_DEFAULT_SENDER'],
            recipients=[user.email]
        )
        msg.body = (
            f"Hello,\n\n"
            f"Your username is: {user.username}\n\n"
            "If you did not request this, secure your account immediately.\n\n"
            "Thank you,\nNoteSaver Pro Team"
        )
        thread = threading.Thread(target=_send_mail_async, args=(app, msg))
        thread.daemon = True
        thread.start()
        return True
    except Exception as e:
        logger.error(f"Error queuing username reminder to {user.email}: {e}")
        return False


        


@app.route('/register', methods=['GET', 'POST'])
def register():
    form = RegistrationForm()
    
    otp_sent_flag = False
    user_email_for_template = None 
    
    # Check session for pre-registered but unverified user email
    # This helps persist the email between the two steps even if the user refreshes
    unverified_email = session.get('unverified_user_email')
    
    if unverified_email:
        otp_sent_flag = True
        user_email_for_template = unverified_email

    if request.method == 'POST':
        action = request.form.get('action')
        
        # Determine the target email for verification steps
        if action in ['verify_otp', 'resend']:
            user_email_for_template = request.form.get('current_user_email')
            # If the user posts Step 2, they are already past Step 1
            if user_email_for_template:
                otp_sent_flag = True

        # --- ACTION 1: REGISTER (Step 1 Submission) ---
        if action == 'register' and form.validate_on_submit():
            
            # Check if user already exists by username or email
            existing_user = User.query.filter(
                (User.username == form.username.data) | 
                (User.email == form.email.data)
            ).first()
            if existing_user:
                flash("An account with this username or email already exists.", "danger")
                return render_template('register.html', form=form, otp_sent=False)

            # Check mobile number duplicate (only if provided)
            mobile = getattr(form, 'mobile_number', None)
            mobile_value = mobile.data.strip() if mobile and mobile.data else None
            if mobile_value:
                existing_mobile = User.query.filter_by(mobile_number=mobile_value).first()
                if existing_mobile:
                    flash("An account with this mobile number already exists.", "danger")
                    return render_template('register.html', form=form, otp_sent=False)

            # ── Check 4: Deleted account — 7-day cooldown ──
            deleted_by_email = DeletedAccount.query.filter_by(email=form.email.data).first()
            if deleted_by_email:
                if deleted_by_email.is_cooldown_active():
                    days_left = deleted_by_email.cooldown_remaining_days()
                    flash(
                        f"This email was recently used for a deleted account. "
                        f"You can re-register after {days_left} more day(s).",
                        "warning"
                    )
                    return render_template('register.html', form=form, otp_sent=False)
                else:
                    # Cooldown khatam — purana record delete karo, fresh start allow karo
                    db.session.delete(deleted_by_email)
                    db.session.commit()

            if mobile_value:
                deleted_by_mobile = DeletedAccount.query.filter_by(mobile_number=mobile_value).first()
                if deleted_by_mobile:
                    if deleted_by_mobile.is_cooldown_active():
                        days_left = deleted_by_mobile.cooldown_remaining_days()
                        flash(
                            f"This mobile number was recently used for a deleted account. "
                            f"You can re-register after {days_left} more day(s).",
                            "warning"
                        )
                        return render_template('register.html', form=form, otp_sent=False)
                    else:
                        # Cooldown khatam — purana record delete karo
                        db.session.delete(deleted_by_mobile)
                        db.session.commit()

            # ── Unverified ghost record cleanup ──
            ghost_user = User.query.filter_by(
                email=form.email.data, is_verified=False
            ).first()
            if ghost_user:
                db.session.delete(ghost_user)
                db.session.commit()

            # ── Naya unverified user banao ──
            user = User(
                username=form.username.data,
                email=form.email.data,
                password_hash=generate_password_hash(form.password.data),
                is_verified=False,
                first_name=(form.first_name.data.strip() if getattr(form, 'first_name', None) and form.first_name.data else None),
                last_name=(form.last_name.data.strip() if getattr(form, 'last_name', None) and form.last_name.data else None),
                mobile_number=mobile_value or None,
            )
            
            # Generate OTP and Expiry
            otp_code = generate_otp()
            user.otp_code = otp_code
            user.otp_expiry = datetime.now() + timedelta(minutes=5)
            
            # Save unverified user to database
            db.session.add(user)
            db.session.commit()

            # SEND EMAIL (The Fix for 'otp nahi aaya')
            if send_verification_email(user, otp_code):
                flash(f"A verification code has been sent to {user.email}. Check your spam folder.", "info")
                # Persist the email in the session
                session['unverified_user_email'] = user.email
                
                return render_template('register.html', 
                    form=form, 
                    otp_sent=True, 
                    current_user_email=user.email)
            else:
                flash("Error sending verification email. Please try again later.", "danger")
                # If email fails, delete the user and return to step 1
                db.session.delete(user)
                db.session.commit()
                return render_template('register.html', form=form, otp_sent=False)

        # --- ACTION 2: VERIFY OTP (Step 2 Submission) ---
        elif action == 'verify_otp':
            otp = request.form.get('otp')
            
            if is_otp_valid(user_email_for_template, otp):
                user = User.query.filter_by(email=user_email_for_template).first()
                if user:
                    user.is_verified = True
                    user.otp_code = None # Clear OTP after verification
                    user.otp_expiry = None
                    db.session.commit()
                    session.pop('unverified_user_email', None) # Clear session data
                    
                    login_user(user)
                    flash("Registration successful! Your email is verified.", "success")
                    return redirect(url_for('dashboard')) 
                
            # If OTP is invalid or user not found
            flash("Invalid or expired OTP. Please try again.", "danger")
            # Re-render step 2 with error
            return render_template('register.html', 
                form=form, 
                otp_sent=True, 
                current_user_email=user_email_for_template)

        # --- ACTION 3: RESEND CODE ---
        elif action == 'resend':
            user = User.query.filter_by(email=user_email_for_template).first()
            if user:
                # Generate new OTP and update DB
                new_otp = generate_otp()
                user.otp_code = new_otp
                user.otp_expiry = datetime.now() + timedelta(minutes=5)
                db.session.commit()
                
                # Resend the email
                send_verification_email(user, new_otp)
                flash("New verification code has been resent to your email.", "info")
                
                return render_template('register.html', 
                    form=form, 
                    otp_sent=True, 
                    current_user_email=user_email_for_template)
            else:
                flash("Could not find account to resend OTP.", "danger")
                return render_template('register.html', form=form, otp_sent=False)


    # --- Final Template Render (GET Request or non-redirect POST) ---
    return render_template(
        'register.html', 
        form=form, 
        otp_sent=otp_sent_flag, 
        current_user_email=user_email_for_template
    )
    
    


@app.route('/request_username', methods=['GET', 'POST'])
@limiter.limit("5 per hour") # Rate limit username requests
def request_username_reminder():
    """Renders a form to request a username reminder via email."""
    form = RequestUsernameForm()
    
    if form.validate_on_submit():
        user = User.query.filter_by(email=form.email.data).first()
        
        # 2. सुरक्षा: एक सामान्य संदेश दिखाएं, भले ही उपयोगकर्ता मौजूद हो या न हो।
        if user:
            send_username_reminder_email(user)
        
        flash('If an account is associated with that email, your username has been sent to your inbox.', 'info')
        return redirect(url_for('login'))
        
    # 3. टेम्पलेट को form ऑब्जेक्ट पास करें
    return render_template('request_username_reminder.html', 
                           title='Forgot Username',
                           form=form) # <-- FIX: form object is now passed
    

@app.route('/login', methods=['GET', 'POST'])
@limiter.limit("10 per minute")
@limiter.limit("50 per hour")
@limiter.limit("200 per day")
def login():

    form = LoginForm()

    if form.validate_on_submit():
        user = User.query.filter_by(username=form.username.data).first()

        if user and user.check_password(form.password.data):

            # ── next page save karo PEHLE session.clear() se ─
            from urllib.parse import urlparse
            next_page = request.args.get('next') or request.form.get('next') or ''
            if next_page and urlparse(next_page).netloc != '':
                next_page = ''  # external URL block karo

            # ── clear old session ─────────────────────────────
            session.clear()
            login_user(user, remember=False)

            # ── detect device ─────────────────────────────────
            ua_string = request.headers.get('User-Agent', '')
            from user_agents import parse as _ua_parse
            ua = _ua_parse(ua_string)

            # Browser + version + OS — e.g. "Chrome 124 on Windows 10"
            browser = ua.browser.family or "Unknown Browser"
            browser_ver = ua.browser.version_string.split('.')[0] if ua.browser.version_string else ""
            os_name = ua.os.family or "Unknown OS"
            os_ver = ua.os.version_string.split('.')[0] if ua.os.version_string else ""

            browser_str = f"{browser} {browser_ver}".strip()
            os_str = f"{os_name} {os_ver}".strip()
            device = f"{browser_str} on {os_str}"

            # ── create unique device session ──────────────────
            session_token = str(uuid.uuid4())

            # ── MAX 2 DEVICES LIMIT ────────────────────────────
            MAX_ACTIVE_SESSIONS = 2

            # Pehle ALL active sessions deactivate karo jo limit se zyada hain
            # created_at se sort — oldest first
            active_sessions = UserSession.query.filter_by(
                user_id=user.id,
                is_active=True
            ).order_by(UserSession.created_at.asc()).all()

            logger.info(f"Login: {user.username} has {len(active_sessions)} active sessions before new login")

            # Agar 2 ya zyada active hain to oldest(s) deactivate karo
            # taaki naye ke saath total 2 rahein
            while len(active_sessions) >= MAX_ACTIVE_SESSIONS:
                oldest = active_sessions.pop(0)  # sabse purani
                oldest.is_active = False
                logger.info(
                    f"Session limit hit — force-logout: {oldest.device} "
                    f"token={oldest.session_token[:8]}"
                )

            db.session.flush()  # deactivation DB mein likho commit se pehle
            # ──────────────────────────────────────────────────

            new_session = UserSession(
                user_id=user.id,
                session_token=session_token,
                device=device,
                ip_address=request.remote_addr,
                location="Detecting...",
                created_at=datetime.now(timezone.utc),
                last_activity=datetime.now(timezone.utc),
                is_active=True
            )

            db.session.add(new_session)
            db.session.commit()

            # save device token
            session['session_token'] = session_token
            session.permanent = False

            flash('Logged in successfully.', 'success')
            if next_page:
                return redirect(next_page)
            return redirect(url_for('dashboard'))

        flash('Invalid username or password.', 'danger')

    return render_template('login.html', form=form)







@app.before_request
def make_session_persistent():
    request.start_time = time.time()
    session.permanent = False  # ← yeh line zaroori hai! Browser close = logout
        


@app.route('/logout')
@login_required
def logout():

    token = session.get('session_token')

    if token:
        db_session = UserSession.query.filter_by(session_token=token).first()
        if db_session:
            db_session.is_active = False
            db_session.last_activity = datetime.now(timezone.utc)
            db.session.commit()

    session.clear()
    logout_user()

    flash('Logged out successfully', 'success')
    return redirect(url_for('index'))



@app.route('/dashboard')
@login_required
@limiter.limit("100 per minute")
def dashboard():
    """
    Dashboard - सिर्फ owner के notes दिखाओ
    Collaborators को dashboard access नहीं है
    """
    page = request.args.get('page', 1, type=int)
    category = request.args.get('category', '')
    search = request.args.get('search', '')

    # सिर्फ अपने notes
    query = Note.query.filter_by(user_id=current_user.id)
    
    if category:
        query = query.filter_by(category=category)
    if search:
        search_pattern = f"%{search}%"
        query = query.filter(
            (Note.title.ilike(search_pattern)) | (Note.content.ilike(search_pattern))
        )

    notes = query.order_by(Note.updated_at.desc()).paginate(
        page=page, per_page=6, error_out=False
    )

    categories = db.session.query(Note.category).filter_by(user_id=current_user.id).distinct().all()
    categories = [c[0] for c in categories]

    return render_template('dashboard.html', notes=notes, categories=categories, 
                         current_category=category, search=search)
    
@app.route('/api/auto-save', methods=['POST'])
@login_required
@limiter.limit("60 per minute")
def auto_save_note():
    """
    Auto-save note draft.
    Storage priority: Redis (fast) → DB Note.draft_content column (fallback) → session (last resort).
    Session is NOT used for content to avoid session bloat / corruption.
    """
    try:
        data = request.get_json()
        if not data:
            return jsonify({'status': 'error', 'message': 'No data received'}), 400

        content = data.get('content', '')
        title   = data.get('title', '')
        note_id = data.get('noteId')

        # Size guard
        if len(str(content)) > MAX_CONTENT_SIZE:
            return jsonify({
                'status': 'error',
                'message': f'Content too large! Maximum {MAX_CONTENT_SIZE:,} characters.'
            }), 400

        if not title and not content:
            return jsonify({'status': 'error', 'message': 'No data to save'}), 400

        draft_payload = json.dumps({
            'title':      title,
            'category':   data.get('category', ''),
            'content':    content,
            'paper_size': data.get('paperSize', 'plain'),
            'note_id':    note_id,
            'timestamp':  datetime.now(timezone.utc).isoformat(),
        }, ensure_ascii=False)

        saved_via = 'none'

        # ── 1. Redis (preferred — fast, no DB load) ───────────────
        if REDIS_AVAILABLE and redis_client:
            try:
                redis_key = f'autosave:{current_user.id}'
                redis_client.setex(redis_key, 7200, draft_payload)  # 2 hr TTL
                saved_via = 'redis'
            except Exception as re:
                logger.warning(f"Redis auto-save failed for {current_user.username}: {re}")

        # ── 2. DB fallback — update existing note draft column ────
        if saved_via == 'none' and note_id:
            try:
                note = Note.query.filter_by(id=note_id, user_id=current_user.id).first()
                if note and hasattr(note, 'draft_content'):
                    note.draft_content = content
                    note.draft_saved_at = datetime.now(timezone.utc)
                    db.session.commit()
                    saved_via = 'db'
            except Exception as de:
                logger.warning(f"DB auto-save fallback failed: {de}")
                db.session.rollback()

        # ── 3. Session last-resort (metadata only, NO content) ────
        if saved_via == 'none':
            session[f'autosave_meta_{current_user.id}'] = {
                'title':     title[:200],
                'note_id':   note_id,
                'timestamp': datetime.now(timezone.utc).isoformat(),
            }
            saved_via = 'session-meta'
            logger.warning(f"Auto-save degraded to session-meta for {current_user.username}")

        logger.info(f"Auto-saved via [{saved_via}] for user {current_user.username}")
        return jsonify({
            'status':    'success',
            'message':   'Auto-saved successfully',
            'saved_via': saved_via,
            'timestamp': datetime.now(timezone.utc).isoformat(),
        })

    except Exception as e:
        logger.error(f"Auto-save critical error for {current_user.username}: {e}")
        return jsonify({'status': 'error', 'message': 'Auto-save failed'}), 500


@app.route('/api/get-auto-save')
@login_required
@limiter.limit("30 per minute")
def get_auto_save():
    """Retrieve auto-saved draft — checks Redis first, then session meta."""
    try:
        # ── Redis ────────────────────────────────────────────────
        if REDIS_AVAILABLE and redis_client:
            try:
                redis_key = f'autosave:{current_user.id}'
                raw = redis_client.get(redis_key)
                if raw:
                    draft = json.loads(raw)
                    saved_time = datetime.fromisoformat(draft['timestamp'].replace('Z', '+00:00'))
                    diff = datetime.now(timezone.utc) - saved_time
                    if diff.total_seconds() < 7200:
                        draft['time_ago'] = _human_time_ago(diff)
                        return jsonify({'status': 'success', 'data': draft})
            except Exception as re:
                logger.warning(f"Redis get-auto-save failed: {re}")

        # ── Session meta fallback ─────────────────────────────────
        meta = session.get(f'autosave_meta_{current_user.id}')
        if meta:
            try:
                saved_time = datetime.fromisoformat(meta['timestamp'].replace('Z', '+00:00'))
                diff = datetime.now(timezone.utc) - saved_time
                if diff.total_seconds() < 3600:
                    meta['time_ago'] = _human_time_ago(diff)
                    meta['content'] = ''   # content not in session
                    return jsonify({'status': 'success', 'data': meta, 'partial': True})
            except Exception:
                pass

        return jsonify({'status': 'no_data'})

    except Exception as e:
        logger.error(f"get-auto-save error: {e}")
        return jsonify({'status': 'error'}), 500


def _human_time_ago(diff: timedelta) -> str:
    """Convert timedelta to human-readable string."""
    secs = int(diff.total_seconds())
    if secs < 60:
        return "just now"
    if secs < 3600:
        return f"{secs // 60}m ago"
    return f"{secs // 3600}h ago"

def is_ajax():
    return request.headers.get('X-Requested-With') == 'XMLHttpRequest'




def is_ajax_request():
    """Check if request is AJAX/JSON"""
    return (
        request.headers.get('X-Requested-With') == 'XMLHttpRequest' or
        'application/json' in request.headers.get('Accept', '') or
        request.is_json
    )

def ajax_error_response(message, status_code=400):
    """Return consistent JSON error response"""
    return jsonify({
        'success': False,
        'message': message,
        'status': status_code
    }), status_code

def ajax_success_response(message, redirect_url=None, **kwargs):
    """Return consistent JSON success response"""
    response = {
        'success': True,
        'message': message,
    }
    if redirect_url:
        response['redirect_url'] = redirect_url
    response.update(kwargs)
    return jsonify(response)


@app.route('/create_note', methods=['GET', 'POST'])
@login_required
@limiter.limit("30 per minute")
@limiter.limit("200 per hour")
def create_note():
    if request.method == 'POST':
        try:
            logger.debug("\n" + "="*60)
            logger.debug("📝 CREATE NOTE REQUEST")
            logger.debug("="*60)
            
            # Get form data
            title = request.form.get('title', '').strip()
            content = request.form.get('content', '').strip()
            category = request.form.get('category', 'General').strip()
            paper_size = request.form.get('paper_size', 'plain').strip()
            
            logger.debug(f"Title: {title[:50]}...")
            logger.debug(f"Content length: {len(content)}")
            logger.debug(f"Category: {category}")
            logger.debug(f"Paper Size: {paper_size}")
            logger.debug("="*60)

            # ===== VALIDATION: Title =====
            if not title:
                logger.warning(f"Note creation failed - missing title for user {current_user.username}")
                return jsonify({'success': False, 'message': '❌ Title is required.'}), 400

            if len(title) > 255:
                return jsonify({'success': False, 'message': '❌ Title must be 255 characters or less.'}), 400

            # ===== VALIDATION: Content =====
            if not content:
                logger.warning(f"Note creation failed - missing content for user {current_user.username}")
                return jsonify({'success': False, 'message': '❌ Content is required.'}), 400

            # ===== VALIDATION: Content Size =====
            if len(content) > MAX_CONTENT_SIZE:
                logger.warning(f"Note creation failed - content too large ({len(content)} chars)")
                return jsonify({
                    'success': False, 
                    'message': f'❌ Content too large! Maximum is {MAX_CONTENT_SIZE:,} characters.'
                }), 400

            # ===== VALIDATION: Paper Size =====
            is_valid, error = validate_paper_size(paper_size)
            if not is_valid:
                logger.warning(f"Invalid paper size: {paper_size}")
                return jsonify({'success': False, 'message': error}), 400

            # ===== STORAGE CHECK =====
            if not current_user.is_storage_available(len(content)):
                logger.warning(f"Storage limit exceeded for user {current_user.username}")
                return jsonify({
                    'success': False,
                    'message': '❌ Storage limit exceeded. Please upgrade or delete some notes.'
                }), 400

            # ===== CLEAN HTML =====
            logger.debug("🧹 Cleaning HTML content...")
            try:
                cleaned_content = clean_html(content)
                logger.info(f"✅ HTML cleaned successfully")
            except Exception as clean_error:
                logger.error(f"HTML cleaning error: {clean_error}")
                return jsonify({'success': False, 'message': '❌ Error processing content'}), 400

            # ===== CREATE NOTE OBJECT =====
            logger.debug("📦 Creating Note object...")
            note = Note(
                title=title,
                content=cleaned_content,
                category=category if category else 'General',
                paper_size=paper_size,
                user_id=current_user.id
            )
            
            # Update content stats
            note.update_content_stats()
            
            logger.info(f"✅ Note created: {note.title}")
            logger.debug(f"   Word count: {note.word_count}")
            logger.debug(f"   Estimated pages: {note.estimated_pages}")

            # ===== SAVE TO DATABASE =====
            logger.debug("💾 Saving to database...")
            db.session.add(note)
            db.session.commit()
            logger.info(f"✅ Note saved with ID: {note.id}")
            
            # Clear auto-save draft (Redis + session)
            if REDIS_AVAILABLE and redis_client:
                try:
                    redis_client.delete(f'autosave:{current_user.id}')
                except Exception:
                    pass
            session.pop(f'autosave_meta_{current_user.id}', None)
            
            logger.info(f"✅ Note created: '{note.title}' (ID: {note.id}, Pages: {note.estimated_pages})")

            response = {
                'success': True,
                'message': '✅ Note created successfully!',
                'id': note.id,
                'redirect_url': url_for('dashboard'),
                'stats': {
                    'word_count': note.word_count,
                    'char_count': note.char_count,
                    'estimated_pages': note.estimated_pages
                }
            }
            logger.debug(f"📤 Response: {response}")
            logger.debug("="*60 + "\n")
            return jsonify(response), 201

        except Exception as e:
            logger.error(f"\n❌ EXCEPTION: {type(e).__name__}: {str(e)}")
            logger.debug("="*60)
            
            db.session.rollback()
            logger.error(f"Error creating note: {str(e)}", exc_info=True)
            
            return jsonify({
                'success': False,
                'message': f'❌ Error creating note: {str(e)}'
            }), 500

    # GET REQUEST — check for a saved draft
    auto_save_data = None
    if REDIS_AVAILABLE and redis_client:
        try:
            raw = redis_client.get(f'autosave:{current_user.id}')
            if raw:
                draft = json.loads(raw)
                saved_time = datetime.fromisoformat(draft['timestamp'].replace('Z', '+00:00'))
                diff = datetime.now(timezone.utc) - saved_time
                if diff.total_seconds() < 7200:
                    draft['time_ago'] = _human_time_ago(diff)
                    auto_save_data = draft
        except Exception:
            pass

    if not auto_save_data:
        meta = session.get(f'autosave_meta_{current_user.id}')
        if meta:
            try:
                saved_time = datetime.fromisoformat(meta['timestamp'].replace('Z', '+00:00'))
                diff = datetime.now(timezone.utc) - saved_time
                if diff.total_seconds() < 3600:
                    meta['time_ago'] = _human_time_ago(diff)
                    auto_save_data = meta
            except Exception:
                pass
    
    return render_template('note.html', 
                         action='Create',
                         auto_save_data=auto_save_data)


@app.route('/view_note/<int:note_id>')
@login_required
@check_collaboration_permission
def view_note(note_id):
    """View a note (respecting collaborator permissions)"""
    note = Note.query.get(note_id)
    if not note:
        flash('Note not found.', 'danger')
        return redirect(url_for('dashboard'))

    # Render view template - collaborators with view/comment/edit/admin can view
    return render_template('view_note.html', note=note)


@app.route('/edit_note/<int:note_id>', methods=['GET', 'POST'])
@login_required
@limiter.limit("50 per minute")
@check_collaboration_permission
def edit_note(note_id):
    note, has_access = check_note_access(note_id)

    if not note:
        flash('Note not found.', 'danger')
        return redirect(url_for('dashboard'))

    if not has_access:
        flash('This note is password protected.', 'warning')
        return redirect(url_for('dashboard'))

    if request.method == 'POST':
        try:
            logger.debug("\n" + "="*60)
            logger.debug("🔧 EDIT NOTE REQUEST")
            logger.debug("="*60)

            # ========== 🔧 IMPORTANT FIX: UNIFIED DATA HANDLING ==========
            if request.is_json:
                data = request.get_json(silent=True) or {}
            else:
                data = request.form.to_dict()

            title = (data.get('title') or '').strip()
            content = (data.get('content') or '').strip()
            category = (data.get('category') or 'General').strip()
            paper_size = (data.get('paper_size') or 'plain').strip()
            # =============================================================

            logger.debug(f"Note ID: {note_id}")
            logger.debug(f"Title: {title[:50]}...")
            logger.debug(f"Content length: {len(content)}")
            logger.debug(f"Paper Size: {paper_size}")
            logger.debug("="*60)

            # ========== VALIDATIONS ==========
            if not title:
                logger.warning(f"Edit failed - missing title for note {note_id}")
                return jsonify({'success': False, 'message': '❌ Title is required.'}), 400

            if not content:
                logger.warning(f"Edit failed - missing content for note {note_id}")
                return jsonify({'success': False, 'message': '❌ Content is required.'}), 400

            if len(content) > MAX_CONTENT_SIZE:
                logger.warning(f"Edit failed - content too large for note {note_id}")
                return jsonify({
                    'success': False,
                    'message': f'❌ Content exceeds maximum size ({MAX_CONTENT_SIZE} chars).'
                }), 400

            # Validate paper size
            is_valid, error = validate_paper_size(paper_size)
            if not is_valid:
                return jsonify({'success': False, 'message': error}), 400

            # ========== CLEAN HTML ==========
            logger.debug("🧹 Cleaning HTML content...")
            try:
                cleaned_content = clean_html(content)
                logger.info("✅ HTML cleaned successfully")
            except Exception as clean_error:
                logger.error(f"HTML cleaning error: {clean_error}")
                return jsonify({'success': False, 'message': '❌ Error processing content'}), 400

            # ========== UPDATE NOTE ==========
            logger.debug("📝 Updating note...")
            old_paper_size = note.paper_size
            old_content_size = len(note.content)

            note.title = title
            note.content = cleaned_content
            note.category = category if category else 'General'
            note.paper_size = paper_size
            note.updated_at = datetime.now(timezone.utc)
            note.content_version += 1

            note.update_content_stats()
            db.session.commit()

            logger.info("✅ Note updated successfully")
            logger.debug(f"   Old size: {old_content_size} → New size: {len(note.content)}")
            logger.debug(f"   Old paper size: {old_paper_size} → New: {note.paper_size}")
            logger.debug(f"   Word count: {note.word_count}")
            logger.debug(f"   Estimated pages: {note.estimated_pages}")
            logger.debug("="*60 + "\n")

            logger.info(
                f"✅ Note edited: '{note.title}' (ID: {note_id}, Pages: {note.estimated_pages})"
            )

            success_msg = '✅ Note updated successfully!'

            # ========== AJAX FRIENDLY RESPONSE (VERY IMPORTANT) ==========
            if request.is_json or request.headers.get('X-Requested-With') == 'XMLHttpRequest':
                return jsonify({
                    'success': True,
                    'message': success_msg,
                    'redirect_url': url_for('dashboard'),
                    'note_id': note.id,
                    'stats': {
                        'word_count': note.word_count,
                        'char_count': note.char_count,
                        'estimated_pages': note.estimated_pages
                    }
                }), 200
            # =============================================================

            flash(success_msg, 'success')
            return redirect(url_for('dashboard'))

        except Exception as e:
            logger.error(f"\n❌ EXCEPTION: {type(e).__name__}: {str(e)}")
            logger.debug("="*60)

            db.session.rollback()
            logger.error(f"Error editing note {note_id}: {str(e)}", exc_info=True)

            if request.is_json:
                return jsonify({
                    'success': False,
                    'message': '❌ Error updating note'
                }), 500

            flash('Error updating note.', 'error')
            return redirect(request.url)

    # ========== GET REQUEST ==========
    form_data = {
        'title': note.title,
        'content': note.content,
        'category': note.category,
        'paper_size': note.paper_size
    }

    return render_template(
        'note.html',
        action='Edit',
        form_data=form_data,
        note_id=note_id
    )


# STEP 7: API endpoint for getting note stats
@app.route('/api/note/<int:note_id>/stats')
@login_required
def get_note_stats(note_id):
    """Get note statistics"""
    try:
        note = Note.query.filter_by(id=note_id, user_id=current_user.id).first()
        
        if not note:
            return jsonify({'success': False, 'message': 'Note not found'}), 404
        
        is_valid, msg, details = note.validate_for_unified_flow()
        
        return jsonify({
            'success': True,
            'stats': {
                'word_count': note.word_count,
                'char_count': note.char_count,
                'estimated_pages': note.estimated_pages,
                'content_size_mb': note.get_content_size_mb(),
                'paper_size': note.paper_size,
                'paper_config': note.get_paper_config(),
                'validation': {
                    'is_valid': is_valid,
                    'message': msg,
                    'details': details
                }
            }
        }), 200
        
    except Exception as e:
        logger.error(f"Error fetching note stats: {e}")
        return jsonify({'success': False, 'message': 'Error loading stats'}), 500


# STEP 8: User storage API
@app.route('/api/user/storage')
@login_required
def get_user_storage():
    """Get user storage information"""
    try:
        used_mb = current_user.get_storage_used_mb()
        max_mb = current_user.get_max_storage_mb()
        used_percent = (used_mb / max_mb * 100) if max_mb > 0 else 0
        
        return jsonify({
            'success': True,
            'storage': {
                'used_mb': used_mb,
                'max_mb': max_mb,
                'used_percent': round(used_percent, 2),
                'available': current_user.is_storage_available(),
                'premium': current_user.is_premium_active(),
                'notes_count': current_user.get_notes_count()
            }
        }), 200
        
    except Exception as e:
        logger.error(f"Error fetching storage info: {e}")
        return jsonify({'success': False, 'message': 'Error loading storage info'}), 500

# Add route for the multi-page editor
@app.route("/multi-page-editor")
@login_required
@limiter.limit("20 per minute")
def multi_page_editor():
    """Direct route to multi-page editor"""
    return render_template("note_editor.html")

@app.route('/delete_note/<int:note_id>', methods=['POST'])
@login_required
@limiter.limit("20 per minute")  # Delete note limit
def delete_note(note_id):
    note, has_access = check_note_access(note_id)
    
    if not note:
        flash('Note not found.', 'danger')
        return redirect(url_for('dashboard'))
    
    if not has_access:
        flash('This note is password protected. Please verify password from dashboard.', 'warning')
        return redirect(url_for('dashboard'))
    
    try:
        note_title = note.title
        db.session.delete(note)
        db.session.commit()
        
        # Clear session verification
        session.pop(f'note_verified_{note_id}', None)
        
        logger.info(f"Note deleted by user {current_user.username}: {note_title}")
        flash('Note deleted successfully!', 'success')
    except Exception as e:
        db.session.rollback()
        logger.error(f"Error deleting note {note_id}: {e}")
        flash('Error deleting note. Please try again.', 'error')
    
    return redirect(url_for('dashboard'))


# ─── NEW: Duplicate a note ───────────────────────────────────────────────────
@app.route('/duplicate_note/<int:note_id>', methods=['POST'])
@login_required
@limiter.limit("20 per minute")
def duplicate_note(note_id):
    """Create a copy of an existing note."""
    note, has_access = check_note_access(note_id)
    if not note:
        return jsonify({'success': False, 'message': 'Note not found.'}), 404
    if not has_access:
        return jsonify({'success': False, 'message': 'Password protected note.'}), 403
    try:
        copy = Note(
            title=f"{note.title} (Copy)",
            content=note.content,
            category=note.category,
            paper_size=note.paper_size,
            user_id=current_user.id,
        )
        copy.update_content_stats()
        db.session.add(copy)
        db.session.commit()
        logger.info(f"Note duplicated: {note.title} → {copy.id} by {current_user.username}")
        return jsonify({
            'success': True,
            'message': '✅ Note duplicated successfully!',
            'new_note_id': copy.id,
            'redirect_url': url_for('dashboard'),
        }), 201
    except Exception as e:
        db.session.rollback()
        logger.error(f"Error duplicating note {note_id}: {e}")
        return jsonify({'success': False, 'message': 'Error duplicating note.'}), 500


@app.route('/toggle_favorite/<int:note_id>', methods=['POST'])
@login_required
@premium_required
@limiter.limit("100 per minute")  # Favorite toggle limit
def toggle_favorite(note_id):
    note = Note.query.get_or_404(note_id)
    if note.user_id != current_user.id:
        return jsonify({'status': 'error', 'message': 'Unauthorized'}), 403

    try:
        note.is_favorite = not note.is_favorite
        db.session.commit()
        
        logger.info(f"Note favorite toggled by user {current_user.username}: {note.title}")
        return jsonify({'status': 'success', 'is_favorite': note.is_favorite})
    except Exception as e:
        db.session.rollback()
        logger.error(f"Error toggling favorite for note {note_id}: {e}")
        return jsonify({'status': 'error', 'message': 'Database error'}), 500

# ... (Previous imports and code remain unchanged until the download_pdf route)

# ============= UNIVERSAL PDF DOWNLOAD (ALL LANGUAGES) =============
# Supports: Hindi, English, Chinese, Arabic, Japanese, Korean, etc.



@app.route('/download/pdf/<int:note_id>')
@login_required
@premium_required
@limiter.limit("10 per minute;50 per hour;200 per day")
def download_pdf(note_id):
    """Generate PDF with multi-language support"""
    # Download options (from query params — set by frontend modal)
    include_title    = request.args.get('include_title',    '1') == '1'
    include_meta     = request.args.get('include_meta',     '1') == '1'
    include_footer   = request.args.get('include_footer',   '1') == '1'

    note, has_access = check_note_access(note_id)
    if not note:
        if is_ajax_request():
            return ajax_error_response('Note not found.', 404)
        flash('Note not found.', 'danger')
        return redirect(url_for('dashboard'))
    
    if not has_access:
        if is_ajax_request():
            return ajax_error_response('Password protected. Verify from dashboard.', 403)
        flash('Verify password from dashboard.', 'warning')
        return redirect(url_for('dashboard'))

    # Redis tracking
    if REDIS_AVAILABLE:
        pdf_key = f"pdf_downloads_{current_user.id}"
        try:
            daily_downloads = int(redis_client.get(pdf_key) or 0)
            max_daily = 50 if current_user.is_premium_active() else 10
            if daily_downloads >= max_daily:
                if is_ajax_request():
                    return ajax_error_response(f'Daily limit ({max_daily}) exceeded!', 429)
                flash(f'Daily PDF limit ({max_daily}) exceeded!', 'warning')
                return redirect(url_for('dashboard'))
            redis_client.incr(pdf_key)
            redis_client.expire(pdf_key, 86400)
        except Exception as e:
            logger.debug(f"Suppressed error: {e}")

    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont
        from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_JUSTIFY
        from reportlab.lib import colors
        
        buffer = BytesIO()
        doc = SimpleDocTemplate(
            buffer,
            pagesize=A4,
            topMargin=0.75*inch,
            bottomMargin=0.75*inch,
            leftMargin=0.75*inch,
            rightMargin=0.75*inch
        )
        
        styles = getSampleStyleSheet()
        story = []

        # ===== MULTI-LANGUAGE FONT REGISTRATION =====
        # This will work for ALL languages including Hindi, Arabic, Chinese, etc.
        
        font_name = 'Helvetica'  # Fallback
        font_registered = False
        
        try:
            # PRIORITY 1: Try Google Noto Sans (Best for all languages)
            noto_paths = [
                # Windows paths
                'C:\\Windows\\Fonts\\NotoSans-Regular.ttf',
                'C:\\Windows\\Fonts\\NotoSansDevanagari-Regular.ttf',
                # Linux paths
                '/usr/share/fonts/truetype/noto/NotoSans-Regular.ttf',
                '/usr/share/fonts/truetype/noto/NotoSansDevanagari-Regular.ttf',
                '/usr/share/fonts/noto/NotoSans-Regular.ttf',
                # Mac paths
                '/System/Library/Fonts/Supplemental/Arial Unicode.ttf',
            ]
            
            for path in noto_paths:
                if os.path.exists(path):
                    try:
                        pdfmetrics.registerFont(TTFont('UniversalFont', path))
                        font_name = 'UniversalFont'
                        font_registered = True
                        logger.info(f"✅ Universal font registered: {path}")
                        break
                    except Exception as e:
                        logger.warning(f"Failed to register {path}: {e}")
                        continue
            
            # PRIORITY 2: Try Windows built-in fonts
            if not font_registered:
                windows_fonts = [
                    ('C:\\Windows\\Fonts\\arial.ttf', 'Arial'),
                    ('C:\\Windows\\Fonts\\mangal.ttf', 'Mangal'),  # Hindi
                    ('C:\\Windows\\Fonts\\msgothic.ttc', 'MSGothic'),  # Japanese
                    ('C:\\Windows\\Fonts\\simsun.ttc', 'SimSun'),  # Chinese
                ]
                
                for path, name in windows_fonts:
                    if os.path.exists(path):
                        try:
                            pdfmetrics.registerFont(TTFont('UniversalFont', path))
                            font_name = 'UniversalFont'
                            font_registered = True
                            logger.info(f"✅ Registered {name} from {path}")
                            break
                        except Exception as e:
                            continue
            
            # PRIORITY 3: Try Linux fonts
            if not font_registered:
                linux_fonts = [
                    '/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf',
                    '/usr/share/fonts/truetype/liberation/LiberationSans-Regular.ttf',
                    '/usr/share/fonts/truetype/noto-cjk/NotoSansCJK-Regular.ttc',
                ]
                
                for path in linux_fonts:
                    if os.path.exists(path):
                        try:
                            pdfmetrics.registerFont(TTFont('UniversalFont', path))
                            font_name = 'UniversalFont'
                            font_registered = True
                            logger.info(f"✅ Registered Linux font: {path}")
                            break
                        except Exception as e:
                            continue
            
            if not font_registered:
                logger.warning("⚠️ No universal font found! Using Helvetica (limited Unicode support)")
                logger.warning("⚠️ Install Google Noto Sans for full language support")
                
        except Exception as e:
            logger.error(f"Font registration error: {e}")

        # Create custom styles with universal font
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Title'],
            fontName=font_name,
            fontSize=18,
            leading=22,
            alignment=TA_CENTER,
            spaceAfter=12,
            wordWrap='CJK'  # Support for Asian languages
        )
        
        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading1'],
            fontName=font_name,
            fontSize=14,
            leading=18,
            spaceAfter=10,
            wordWrap='CJK'
        )
        
        normal_style = ParagraphStyle(
            'CustomNormal',
            parent=styles['Normal'],
            fontName=font_name,
            fontSize=11,
            leading=16,
            alignment=TA_JUSTIFY,
            wordWrap='CJK'
        )
        
        meta_style = ParagraphStyle(
            'CustomMeta',
            parent=styles['Normal'],
            fontName=font_name,
            fontSize=9,
            textColor=colors.grey,
            alignment=TA_CENTER
        )

        # ── Conditionally add title ──────────────────────────────
        if include_title:
            title_text = note.title or "Untitled Note"
            title_text = title_text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
            story.append(Paragraph(f"<b>{title_text}</b>", title_style))
            story.append(Spacer(1, 0.2*inch))

        # ── Conditionally add metadata ────────────────────────────
        if include_meta:
            meta_text  = f"Category: {note.category or 'General'} | "
            meta_text += f"Created: {note.created_at.strftime('%B %d, %Y')} | "
            meta_text += f"Updated: {note.updated_at.strftime('%B %d, %Y')}"
            story.append(Paragraph(meta_text, meta_style))
            story.append(Spacer(1, 0.3*inch))

        # Process content with proper encoding
        content = note.content.strip() if note.content else ""
        
        if not content or content in ["<p><br></p>", "<p></p>", "<br>", ""]:
            story.append(Paragraph("<i>No content available</i>", normal_style))
        else:
            try:
                # Parse HTML
                soup = BeautifulSoup(content, 'html.parser')
                
                # Remove dangerous tags
                for tag in soup.find_all(['script', 'style', 'iframe', 'object', 'embed', 'meta', 'link']):
                    tag.decompose()

                body = soup.find('body') or soup

                # Process all elements
                for elem in body.find_all(['p', 'div', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'blockquote', 'ul', 'ol', 'pre']):
                    try:
                        # Get text with proper encoding
                        text = elem.get_text(separator=' ', strip=True)
                        
                        if not text or len(text) < 2:
                            continue

                        # Clean text - remove invisible characters
                        text = text.replace('\xa0', ' ')  # Non-breaking space
                        text = text.replace('\u200b', '')  # Zero-width space
                        text = text.replace('\u200c', '')  # Zero-width non-joiner
                        text = text.replace('\u200d', '')  # Zero-width joiner
                        text = ' '.join(text.split())  # Normalize whitespace

                        # Escape XML special characters for ReportLab
                        text = text.replace('&', '&amp;')
                        text = text.replace('<', '&lt;')
                        text = text.replace('>', '&gt;')

                        # Handle different element types
                        if elem.name and elem.name.startswith('h'):
                            # Headings
                            story.append(Paragraph(f"<b>{text}</b>", heading_style))
                            story.append(Spacer(1, 0.1*inch))
                        
                        elif elem.name in ['ul', 'ol']:
                            # Lists
                            for idx, li in enumerate(elem.find_all('li', recursive=False), 1):
                                li_text = li.get_text(separator=' ', strip=True)
                                if li_text:
                                    # Clean and escape
                                    li_text = li_text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                                    bullet = '&#8226;' if elem.name == 'ul' else f'{idx}.'
                                    story.append(Paragraph(f"{bullet} {li_text}", normal_style))
                            story.append(Spacer(1, 0.1*inch))
                        
                        elif elem.name == 'blockquote':
                            # Blockquotes
                            story.append(Paragraph(f'<i>{text}</i>', normal_style))
                            story.append(Spacer(1, 0.1*inch))
                        
                        elif elem.name == 'pre':
                            # Preformatted text (code blocks)
                            from reportlab.platypus import Preformatted
                            story.append(Preformatted(text, styles['Code']))
                            story.append(Spacer(1, 0.1*inch))
                        
                        else:
                            # Regular paragraphs
                            if text:
                                story.append(Paragraph(text, normal_style))
                                story.append(Spacer(1, 0.08*inch))

                    except Exception as elem_error:
                        logger.warning(f"Element processing error: {elem_error}")
                        # Try to add as plain text
                        try:
                            plain_text = elem.get_text(strip=True)
                            if plain_text:
                                plain_text = plain_text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                                story.append(Paragraph(plain_text, normal_style))
                        except Exception as e:
                            continue

                # Handle tables separately
                for table in body.find_all('table'):
                    try:
                        from reportlab.platypus import Table, TableStyle
                        
                        table_data = []
                        for tr in table.find_all('tr'):
                            row = []
                            for cell in tr.find_all(['td', 'th']):
                                cell_text = cell.get_text(strip=True) or ''
                                # Encode properly for all languages
                                cell_text = cell_text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                                row.append(cell_text)
                            if row:
                                table_data.append(row)
                        
                        if table_data and len(table_data) > 0:
                            # Calculate column widths
                            max_cols = max(len(row) for row in table_data)
                            # Ensure all rows have same column count
                            for row in table_data:
                                while len(row) < max_cols:
                                    row.append('')
                            
                            col_widths = [5.5*inch / max_cols] * max_cols
                            
                            pdf_table = Table(table_data, colWidths=col_widths)
                            pdf_table.setStyle(TableStyle([
                                ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                                ('FONTNAME', (0, 0), (-1, -1), font_name),
                                ('FONTSIZE', (0, 0), (-1, -1), 10),
                                ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
                                ('TOPPADDING', (0, 0), (-1, -1), 8),
                                ('GRID', (0, 0), (-1, -1), 0.5, colors.black),
                                ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                                ('WORDWRAP', (0, 0), (-1, -1), True),
                            ]))
                            
                            story.append(pdf_table)
                            story.append(Spacer(1, 0.2*inch))
                    
                    except Exception as table_error:
                        logger.warning(f"Table processing error: {table_error}")
                        pass

            except Exception as parse_error:
                logger.error(f"Content parsing error: {parse_error}")
                # Fallback: use plain text
                try:
                    plain = get_note_preview(note.content, max_length=10000)
                    plain = plain.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                    story.append(Paragraph(plain, normal_style))
                except Exception as fallback_error:
                    logger.error(f"Fallback error: {fallback_error}")
                    story.append(Paragraph("Error processing content", normal_style))



        # ── Conditionally add footer ──────────────────────────────
        if include_footer:
            story.append(Spacer(1, 0.3*inch))
            footer_text = f"<i>Generated by NoteSaver Pro - {datetime.now().strftime('%B %d, %Y')}</i>"
            story.append(Paragraph(footer_text, meta_style))

        # Build PDF
        doc.build(story)
        buffer.seek(0)

        # Create safe filename
        safe_title = safe_filename(note.title or "note")[:100]
        logger.info(f"✅ PDF generated: {safe_title}.pdf ({len(buffer.getvalue())} bytes)")

        return send_file(
            buffer,
            as_attachment=True,
            download_name=f"{safe_title}.pdf",
            mimetype='application/pdf'
        )

    except Exception as e:
        logger.error(f"❌ PDF generation error: {e}")
        logger.error(f"Traceback: {traceback.format_exc()}")
        
        if is_ajax_request():
            return ajax_error_response('PDF generation failed. Try DOCX format.', 500)
        flash('PDF generation failed. Try DOCX format.', 'error')
        return redirect(url_for('dashboard'))


# ============= HELPER FUNCTION =============
def safe_filename(filename):
    """Create safe filename by removing special characters"""
    if not filename:
        return "note"
    # Remove special characters
    safe = re.sub(r'[^a-zA-Z0-9_\-\.]', '_', filename)
    # Remove multiple underscores
    safe = re.sub(r'_+', '_', safe)
    # Trim length
    return safe[:100] if len(safe) > 100 else safe

# ============= DOWNLOAD DOC ROUTE WITH HINDI SUPPORT =============
@app.route('/download/doc/<int:note_id>')
@login_required
@limiter.limit("10 per minute;50 per hour;200 per day")
def download_doc(note_id):
    """Generate DOCX with full Hindi/Devanagari support"""
    # Download options (from query params — set by frontend modal)
    include_title    = request.args.get('include_title',    '1') == '1'
    include_meta     = request.args.get('include_meta',     '1') == '1'
    include_footer   = request.args.get('include_footer',   '1') == '1'

    note, has_access = check_note_access(note_id)
    if not note:
        logger.warning(f"DOC download attempt on invalid note {note_id} by {current_user.username}")
        if is_ajax_request():
            return ajax_error_response('Note not found.', 404, error_code='NOTE_NOT_FOUND')
        flash('Note not found.', 'danger')
        return redirect(url_for('dashboard'))
    
    if not has_access:
        logger.warning(f"DOC download access denied for note {note_id} by {current_user.username}")
        if is_ajax_request():
            return ajax_error_response('Password verification required to download.', 403, error_code='PASSWORD_REQUIRED')
        flash('Verify password from dashboard.', 'warning')
        return redirect(url_for('dashboard'))

    try:
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from bs4 import BeautifulSoup
        
        # Create document
        doc = Document()
        
        # Set margins
        for section in doc.sections:
            section.top_margin = Inches(1.0)
            section.bottom_margin = Inches(1.0)
            section.left_margin = Inches(1.0)
            section.right_margin = Inches(1.0)
        
        # Function to set Hindi font at XML level
        def set_hindi_font_for_run(run, font_name='Mangal', font_size=11, bold=False, italic=False):
            """
            Set Hindi font at the deepest XML level - this is the KEY fix
            """
            # Set basic properties
            run.font.name = font_name
            run.font.size = Pt(font_size)
            run.bold = bold
            run.italic = italic
            
            # Get the run's XML element
            r = run._element
            rPr = r.get_or_add_rPr()
            
            # Remove any existing font elements
            for old_font in rPr.findall(qn('w:rFonts')):
                rPr.remove(old_font)
            
            # Create new font element with ALL font variants
            rFonts = OxmlElement('w:rFonts')
            rFonts.set(qn('w:ascii'), font_name)        # For ASCII chars
            rFonts.set(qn('w:hAnsi'), font_name)        # For high ANSI
            rFonts.set(qn('w:cs'), font_name)           # For complex scripts (CRITICAL for Hindi)
            rFonts.set(qn('w:eastAsia'), font_name)     # For East Asian
            
            # Insert at the beginning of run properties
            rPr.insert(0, rFonts)
            
            return run
        
        def add_paragraph_with_hindi(doc, text, alignment=WD_PARAGRAPH_ALIGNMENT.LEFT, 
                                    font_size=11, bold=False, italic=False, is_heading=False):
            """
            Add paragraph with proper Hindi font support
            """
            if is_heading:
                p = doc.add_heading('', level=2)
                font_size = 14
                bold = True
            else:
                p = doc.add_paragraph()
            
            # Add text as a run
            run = p.add_run(text)
            
            # Apply Hindi font
            set_hindi_font_for_run(run, 'Mangal', font_size, bold, italic)
            
            # Set alignment
            p.alignment = alignment
            
            return p
        
        # Set document properties
        doc.core_properties.title = note.title or "Untitled"
        doc.core_properties.author = current_user.username



        # ── Conditionally add title ──────────────────────────────
        if include_title:
            title_text = note.title or "Untitled"
            title_p = doc.add_heading('', level=1)
            title_run = title_p.add_run(title_text)
            set_hindi_font_for_run(title_run, 'Mangal', 18, bold=True)
            title_p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # ── Conditionally add metadata ────────────────────────────
        if include_meta:
            meta_text  = f"Category: {note.category or 'General'} | "
            meta_text += f"Created: {note.created_at.strftime('%d %b %Y')} | "
            meta_text += f"Updated: {note.updated_at.strftime('%d %b %Y')}"
            meta_p = doc.add_paragraph()
            meta_run = meta_p.add_run(meta_text)
            set_hindi_font_for_run(meta_run, 'Mangal', 9, italic=True)
            meta_p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Add spacing
        doc.add_paragraph()

        # Process content
        content = note.content.strip() if note.content else ""
        
        if not content or content in ["<p><br></p>", "<p></p>", "<br>", ""]:
            p = doc.add_paragraph()
            run = p.add_run("No content available")
            set_hindi_font_for_run(run, 'Mangal', 11, italic=True)
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        else:
            # Parse HTML
            soup = BeautifulSoup(content, 'html.parser')
            
            # Remove unwanted tags
            for tag in soup.find_all(['script', 'style', 'iframe', 'object', 'embed']):
                tag.decompose()
            
            # Get structured lines
            try:
                lines = html_to_plain_lines(content)
            except NameError:
                # Fallback if html_to_plain_lines is not defined
                lines = []
                text = soup.get_text(separator='\n', strip=True)
                for line in text.split('\n'):
                    line = line.strip()
                    if line:
                        lines.append({'text': line, 'align': 'left', 'bold': False, 'italic': False, 'is_heading': False})
            
            if not lines:
                # Fallback: plain text
                text = soup.get_text(separator='\n', strip=True)
                for line in text.split('\n'):
                    line = line.strip()
                    if line:
                        add_paragraph_with_hindi(doc, line)
            else:
                # Process structured content
                for line_info in lines:
                    text = line_info.get('text', '').strip()
                    if not text:
                        continue
                    
                    # Get alignment
                    align_map = {
                        'left': WD_PARAGRAPH_ALIGNMENT.LEFT,
                        'center': WD_PARAGRAPH_ALIGNMENT.CENTER,
                        'right': WD_PARAGRAPH_ALIGNMENT.RIGHT,
                        'justify': WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                    }
                    alignment = align_map.get(line_info.get('align', 'left'), WD_PARAGRAPH_ALIGNMENT.LEFT)
                    
                    # Add paragraph with proper formatting
                    add_paragraph_with_hindi(
                        doc, 
                        text,
                        alignment=alignment,
                        bold=line_info.get('bold', False),
                        italic=line_info.get('italic', False),
                        is_heading=line_info.get('is_heading', False)
                    )



        # ── Conditionally add footer ──────────────────────────────
        if include_footer:
            doc.add_paragraph()
            footer_text = f"Generated by NoteSaver Pro - {datetime.now().strftime('%d %B %Y')}"
            footer_p = doc.add_paragraph()
            footer_run = footer_p.add_run(footer_text)
            set_hindi_font_for_run(footer_run, 'Mangal', 9, italic=True)
            footer_p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Save to buffer
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        safe_title = safe_filename(note.title or "note")
        logger.info(f"✅ DOCX with Hindi support: {safe_title}.docx")

        return send_file(
            buffer,
            as_attachment=True,
            download_name=f"{safe_title}.docx",
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )

    except Exception as e:
        logger.error(f"❌ DOCX error: {e}")
        logger.error(traceback.format_exc())
        if is_ajax_request():
            return ajax_error_response('DOCX generation failed.', 500, error_code='SERVER_ERROR')
        flash('DOCX generation failed.', 'error')
        return redirect(url_for('dashboard'))

# ============= DOWNLOAD TXT ROUTE =============
@app.route('/download/txt/<int:note_id>')
@login_required
@limiter.limit("10 per minute;50 per hour;200 per day")
def download_txt(note_id):
    note, has_access = check_note_access(note_id)
    if not note:
        logger.warning(f"TXT download attempt on invalid note {note_id} by {current_user.username}")
        if is_ajax_request():
            return ajax_error_response('Note not found.', 404, error_code='NOTE_NOT_FOUND')
        flash('Note not found.', 'danger')
        return redirect(url_for('dashboard'))
    
    if not has_access:
        logger.warning(f"TXT download access denied for note {note_id} by {current_user.username}")
        if is_ajax_request():
            return ajax_error_response('Password verification required to download.', 403, error_code='PASSWORD_REQUIRED')
        flash('Verify password to download.', 'warning')
        return redirect(url_for('dashboard'))

    try:
        title = note.title or "Untitled Note"
        soup = BeautifulSoup(note.content, 'html.parser')
        text_content = soup.get_text(separator='\n', strip=True)
        content = f"{title}\n\n{text_content}"
        
        buffer = io.StringIO()
        buffer.write(content)
        buffer.seek(0)
        
        filename = safe_filename(f"{title}.txt")
        logger.info(f"TXT downloaded for note {note_id} by {current_user.username}")
        return send_file(
            io.BytesIO(buffer.getvalue().encode('utf-8')),
            as_attachment=True,
            download_name=filename,
            mimetype='text/plain'
        )
    except Exception as e:
        logger.error(f"Error generating TXT for note {note_id}: {e}")
        if is_ajax_request():
            return ajax_error_response('Error generating text file.', 500, error_code='SERVER_ERROR')
        flash('Error generating text file.', 'error')
        return redirect(url_for('dashboard'))

# ============= SHARE NOTE ROUTE =============
@app.route('/share/note/<int:note_id>')
@login_required
@limiter.limit("50 per minute")
def share_note(note_id):
    note, has_access = check_note_access(note_id)
    if not note:
        logger.warning(f"Share attempt on invalid note {note_id} by {current_user.username}")
        if is_ajax_request():
            return ajax_error_response('Note not found.', 404, error_code='NOTE_NOT_FOUND')
        flash('Note not found.', 'danger')
        return redirect(url_for('dashboard'))
    
    if not has_access:
        logger.warning(f"Share access denied for note {note_id} by {current_user.username}")
        if is_ajax_request():
            return ajax_error_response('Password verification required to share.', 403, error_code='PASSWORD_REQUIRED')
        flash('Verify password to share.', 'warning')
        return redirect(url_for('dashboard'))

    try:
        share_token = serializer.dumps({'note_id': note_id, 'user_id': current_user.id}, salt='share-salt')
        share_url = url_for('public_share_view', token=share_token, _external=True)
        preview = get_note_preview(note.content, 500)
        logger.info(f"Share link generated for note {note_id} by {current_user.username}")
        if is_ajax_request():
            return ajax_success_response(
                'Share link generated.',
                share_url=share_url,
                preview_content=preview,
                note_id=note_id,
                expires_in="7 days"
            )
        return render_template('share_note.html', note=note, preview_content=preview, share_url=share_url, expires_in="7 days")
    except Exception as e:
        logger.error(f"Error generating share link for note {note_id}: {e}")
        if is_ajax_request():
            return ajax_error_response('Error generating share link.', 500, error_code='SERVER_ERROR')
        flash('Error generating share link.', 'error')
        return redirect(url_for('dashboard'))

# ============= NOTE PREVIEW ROUTE =============
@app.route('/note_preview/<int:note_id>')
@login_required
@limiter.limit("50 per minute")
def note_preview(note_id):
    note, has_access = check_note_access(note_id)
    if not note:
        logger.warning(f"Preview attempt on invalid note {note_id} by {current_user.username}")
        return ajax_error_response('Note not found.', 404, error_code='NOTE_NOT_FOUND')
    
    if not has_access:
        logger.warning(f"Preview access denied for note {note_id} by {current_user.username}")
        return ajax_error_response('Password verification required to access preview.', 403, error_code='PASSWORD_REQUIRED')

    try:
        preview = get_note_preview(note.content, 500)
        logger.info(f"Preview generated for note {note_id} by {current_user.username}")
        return ajax_success_response(
            'Preview generated.',
            preview_content=preview,
            note_id=note_id,
            title=note.title
        )
    except Exception as e:
        logger.error(f"Error generating preview for note {note_id}: {e}")
        return ajax_error_response('Error generating preview.', 500, error_code='SERVER_ERROR')

# ------------------- PASSWORD PROTECTION ROUTES ---------------------

@app.route('/verify_note_password', methods=['POST'])
@limiter.limit("20 per minute")  # Password verification limit
@limiter.limit("100 per hour")   # Hourly limit
@login_required
def verify_note_password():
    """Verify note-specific password for protected notes"""
    try:
        data = request.get_json()
        password = data.get('password')
        note_id = data.get('note_id')
        
        logger.info(f"Password verification attempt for note {note_id} by user {current_user.username}")
        
        if not password or not note_id:
            return jsonify({'status': 'error', 'message': 'Password and note ID required'})
        
        # Get the specific note
        note = Note.query.filter_by(id=note_id, user_id=current_user.id).first()
        if not note:
            logger.warning(f"Note {note_id} not found for user {current_user.username}")
            return jsonify({'status': 'error', 'message': 'Note not found'})
        
        # Check if note is password protected
        if not note.is_private:
            return jsonify({'status': 'error', 'message': 'Note is not password protected'})
        
        # Verify the note-specific password using the model's method
        password_valid = note.check_private_password(password)
        
        if password_valid:
            # Store verification in session for this specific note
            session[f'note_verified_{note_id}'] = True
            logger.info(f"Password verified for note {note_id} by user {current_user.username}")
            return jsonify({'status': 'success', 'message': 'Password verified'})
        else:
            logger.warning(f"Incorrect password for note {note_id} by user {current_user.username}")
            return jsonify({'status': 'error', 'message': 'Incorrect password'})
            
    except Exception as e:
        logger.error(f"Error in verify_note_password: {e}")
        return jsonify({'status': 'error', 'message': 'An error occurred'})

@app.route('/set_note_password/<int:note_id>', methods=['POST'])
@login_required
@limiter.limit("10 per minute")  # Password setting limit
def set_note_password(note_id):
    """Set password protection for a note"""
    try:
        data = request.get_json()
        password = data.get('password')
        
        if not password or len(password) < 6:
            return jsonify({'status': 'error', 'message': 'Password must be at least 6 characters'})
        
        note = Note.query.filter_by(id=note_id, user_id=current_user.id).first()
        if not note:
            return jsonify({'status': 'error', 'message': 'Note not found'})
        
        # Set the password protection using the model's method
        note.set_private_password(password)
        note.is_private = True
        db.session.commit()
        
        logger.info(f"Password protection set for note {note_id} by user {current_user.username}")
        return jsonify({'status': 'success', 'message': 'Password protection enabled'})
        
    except Exception as e:
        db.session.rollback()
        logger.error(f"Error setting note password: {e}")
        return jsonify({'status': 'error', 'message': 'An error occurred'})

@app.route('/remove_note_password/<int:note_id>', methods=['POST'])
@login_required
@limiter.limit("10 per minute")  # Password removal limit
def remove_note_password(note_id):
    """Remove password protection from a note"""
    try:
        data = request.get_json()
        password = data.get('password')
        
        if not password:
            return jsonify({'status': 'error', 'message': 'Password required'})
        
        note = Note.query.filter_by(id=note_id, user_id=current_user.id).first()
        if not note:
            return jsonify({'status': 'error', 'message': 'Note not found'})
        
        # Verify current password using the model's method
        password_valid = note.check_private_password(password)
        
        if not password_valid:
            return jsonify({'status': 'error', 'message': 'Incorrect password'})
        
        # Remove password protection
        note.private_password_hash = None
        note.is_private = False
        db.session.commit()
        
        # Clear session verification
        session.pop(f'note_verified_{note_id}', None)
        
        logger.info(f"Password protection removed for note {note_id} by user {current_user.username}")
        return jsonify({'status': 'success', 'message': 'Password protection removed'})
        
    except Exception as e:
        db.session.rollback()
        logger.error(f"Error removing note password: {e}")
        return jsonify({'status': 'error', 'message': 'An error occurred'})


@app.route('/toggle_private/<int:note_id>', methods=['POST'])
@login_required
@limiter.limit("20 per minute")  # Privacy toggle limit
def toggle_private(note_id):
    try:
        note = Note.query.filter_by(id=note_id, user_id=current_user.id).first()
        if not note:
            return jsonify({'status': 'error', 'message': 'Note not found'}), 404

        # Toggle the private flag
        note.is_private = not note.is_private
        
        # If turning private off, also remove password
        if not note.is_private:
            note.private_password_hash = None
            # Clear session verification
            session.pop(f'note_verified_{note_id}', None)

        db.session.commit()
        
        logger.info(f"Privacy toggled for note {note_id} by user {current_user.username}: {'private' if note.is_private else 'public'}")
        
        return jsonify({
            'status': 'success',
            'is_private': note.is_private,
            'message': f'Note is now {"private" if note.is_private else "public"}'
        })
        
    except Exception as e:
        db.session.rollback()
        logger.error(f"Error toggling privacy for note {note_id}: {e}")
        return jsonify({'status': 'error', 'message': str(e)}), 500

# ------------------- OTHER ROUTES ---------------------

@app.route('/premium')
@login_required
def premium():
    return render_template('premium.html')

@app.route('/upgrade_premium')
@login_required
@limiter.limit("5 per hour")
def upgrade_premium():
    """This route is for development/testing only. In production, use Razorpay."""
    if not app.debug:
        flash('Please use the payment system to upgrade.', 'warning')
        return redirect(url_for('premium'))
    current_user.is_premium = True
    current_user.premium_expiry = datetime.now(timezone.utc) + timedelta(days=30)
    db.session.commit()
    logger.warning(f"[DEV] Premium activated directly for user {current_user.username}")
    flash('Premium membership activated for 30 days! (Development Mode)', 'success')
    return redirect(url_for('dashboard'))

@app.route('/activate-free-premium', methods=['POST'])
@login_required
@limiter.limit("3 per hour")
def activate_free_premium():
    """Free premium activation — limited time offer."""
    try:
        if current_user.is_premium_active():
            flash('Premium is already active on your account!', 'info')
            return redirect(url_for('premium'))

        current_user.is_premium = True
        current_user.premium_expiry = datetime.now(timezone.utc) + timedelta(days=30)
        db.session.commit()

        logger.info(f"Free premium activated for user {current_user.username} (id={current_user.id})")
        flash('\U0001f389 Premium activated successfully! Enjoy all features for 30 days.', 'success')
        return redirect(url_for('dashboard'))

    except Exception as e:
        db.session.rollback()
        logger.error(f"Error activating free premium for {current_user.username}: {e}")
        flash('Something went wrong. Please try again.', 'danger')
        return redirect(url_for('premium'))

@app.route('/api/notes')
@login_required
@premium_required
@limiter.limit("100 per minute")
def api_notes():
    """Premium API: list notes with pagination and optional category filter."""
    page     = request.args.get('page', 1, type=int)
    per_page = min(request.args.get('per_page', 20, type=int), 100)  # max 100
    category = request.args.get('category', '').strip()

    query = Note.query.filter_by(user_id=current_user.id)
    if category:
        query = query.filter_by(category=category)

    paginated = query.order_by(Note.updated_at.desc()).paginate(
        page=page, per_page=per_page, error_out=False
    )

    notes_data = [{
        'id': n.id,
        'title': n.title,
        'category': n.category,
        'paper_size': n.paper_size,
        'is_favorite': n.is_favorite,
        'is_private': n.is_private,
        'word_count': getattr(n, 'word_count', None),
        'created_at': n.created_at.isoformat(),
        'updated_at': n.updated_at.isoformat(),
    } for n in paginated.items]

    logger.info(f"API notes accessed by user {current_user.username}: page={page}")
    return jsonify({
        'success': True,
        'notes': notes_data,
        'pagination': {
            'page': page,
            'per_page': per_page,
            'total': paginated.total,
            'pages': paginated.pages,
            'has_next': paginated.has_next,
            'has_prev': paginated.has_prev,
        }
    })

@app.route('/private_note/create', methods=['GET', 'POST'])
@login_required
@limiter.limit("20 per minute")  # Private note creation limit
def create_private_note():
    form = NoteForm()
    if form.validate_on_submit():
        note = Note(
            title=form.title.data,
            content=clean_html(form.content.data),
            category=form.category.data or 'General',
            user_id=current_user.id,
            is_private=True
        )
        
        # Set password hash if password is provided
        password = form.private_password.data if hasattr(form, 'private_password') else None
        if password:
            note.set_private_password(password)
        
        db.session.add(note)
        db.session.commit()
        
        logger.info(f"Private note created by user {current_user.username}: {note.title}")
        flash('Private note created successfully!', 'success')
        return redirect(url_for('dashboard'))
    return render_template('private_note.html', form=form, action='Create Private Note')

@app.route('/private_note/edit/<int:note_id>', methods=['GET', 'POST'])
@login_required
@limiter.limit("30 per minute")  # Private note editing limit
def edit_private_note(note_id):
    note = Note.query.get_or_404(note_id)
    if note.user_id != current_user.id or not note.is_private:
        flash('Unauthorized access or not a private note!', 'danger')
        return redirect(url_for('dashboard'))

    form = NoteForm()

    if request.method == 'GET':
        form.title.data = note.title
        form.content.data = note.content
        form.category.data = note.category
        if hasattr(form, 'private_password') and hasattr(note, 'private_password'):
            form.private_password.data = note.private_password

    if form.validate_on_submit():
        note.title = form.title.data
        note.content = clean_html(form.content.data)
        note.category = form.category.data or 'General'
        if hasattr(form, 'private_password') and hasattr(note, 'private_password'):
            note.private_password = form.private_password.data
        note.updated_at = datetime.now(timezone.utc)
        db.session.commit()
        
        logger.info(f"Private note edited by user {current_user.username}: {note.title}")
        flash('Private note updated successfully!', 'success')
        return redirect(url_for('dashboard'))

    return render_template('private_note.html', form=form, action='Edit Private Note')

# --------------- Password Reset Routes ------------------

@app.route('/forgot-password', methods=['GET', 'POST'])
@limiter.limit("3 per hour")     # Password reset requests ko strict rakhte hain
@limiter.limit("10 per day")     # Daily limit
def forgot_password():
    form = RequestResetForm()
    if form.validate_on_submit():
        user = User.query.filter_by(email=form.email.data).first()
        if user:
            send_reset_email(user)
            logger.info(f"Password reset requested for email: {form.email.data}")
        else:
            logger.warning(f"Password reset requested for non-existent email: {form.email.data}")
            flash("If your email exists in our system, you will receive a reset link.", "info")
            return redirect(url_for('login'))
        return redirect(url_for('login'))
    return render_template('forgot_password.html', form=form)

@app.route('/reset-password/<token>', methods=['GET', 'POST'])
@limiter.limit("5 per hour")  # Token-based password reset limit
def reset_password(token):
    try:
        email = serializer.loads(token, salt='password-reset-salt', max_age=3600)
    except SignatureExpired:
        flash('The password reset link has expired.', 'danger')
        return redirect(url_for('forgot_password'))
    except BadSignature:
        flash('Invalid password reset token.', 'danger')
        return redirect(url_for('forgot_password'))

    user = User.query.filter_by(email=email).first_or_404()
    form = ResetPasswordForm()

    if form.validate_on_submit():
        user.set_password(form.password.data)
        db.session.commit()
        
        logger.info(f"Password reset completed for user: {user.username}")
        flash('Your password has been updated!', 'success')
        return redirect(url_for('login'))

    return render_template('reset_password.html', form=form)

@app.route('/profile')
@login_required
@limiter.limit("50 per minute")
def profile():
    """
    User profile - सिर्फ अपना profile access कर सकते हो
    """
    # दूसरे के profile देखने की कोशिश?
    requested_user_id = request.args.get('user_id', current_user.id, type=int)
    
    if requested_user_id != current_user.id:
        logger.warning(f"Unauthorized profile access: User {current_user.id} to User {requested_user_id}")
        flash('❌ दूसरे के profile को access नहीं कर सकते', 'danger')
        return redirect(url_for('dashboard'))
    
    return render_template('profile.html')

@app.route('/settings')
@login_required
@limiter.limit("30 per minute")
def settings():
    """
    Settings - सिर्फ अपनी settings access कर सकते हो
    """
    return render_template('settings.html')

# ===== ADD THESE ROUTES TO YOUR app.py =====
# Place these routes around line 1050 (after your existing profile routes)
# ================= REAL LOCATION CAPTURE =================
@app.route("/api/update-location", methods=["POST"])
@login_required
@csrf.exempt
def update_location():

    data = request.get_json(silent=True)

    if not data:
        return jsonify({"success": False, "message": "Invalid JSON"}), 400

    lat = data.get("lat")
    lon = data.get("lon")

    if not lat or not lon:
        return jsonify({"success": False, "message": "Missing coordinates"}), 400

    # Reverse geocoding (OpenStreetMap free API)
    try:
        geo = requests.get(
            f"https://nominatim.openstreetmap.org/reverse?lat={lat}&lon={lon}&format=json",
            headers={"User-Agent": "NoteSaverPro"}
        ).json()

        address = geo.get("address", {})

        city = address.get("city") or address.get("town") or address.get("village")
        state = address.get("state")
        country = address.get("country_code", "").upper()

        location_text = f"{city}, {state}, {country}"

    except Exception as e:
        logger.error("Geo error:", e)
        location_text = "Unknown Location"

    # UPDATE CURRENT SESSION ONLY
    token = session.get("session_token")
    if not token:
        return jsonify({"success": False}), 400

    user_session = UserSession.query.filter_by(
        user_id=current_user.id,
        session_token=token,
        is_active=True
    ).first()

    if user_session:
        user_session.location = location_text
        db.session.commit()

    return jsonify({"success": True, "location": location_text})





@app.route('/api/active-sessions')
@login_required
def active_sessions():

    sessions = UserSession.query.filter_by(
        user_id=current_user.id,
        is_active=True
    ).all()

    data = []
    current_token = session.get('session_token')

    for s in sessions:
        data.append({
            "id": s.id,
            "device": s.device,
            "location": s.location,
            "last_activity": s.last_activity.strftime("%d %b %Y %I:%M %p"),
            "current": s.session_token == current_token
        })

    return jsonify(data)




@app.route('/api/profile/stats')
@login_required
@limiter.limit("30 per minute")
def api_profile_stats():
    try:
        total_notes = Note.query.filter_by(user_id=current_user.id).count()
        last_note = Note.query.filter_by(user_id=current_user.id)\
            .order_by(Note.updated_at.desc()).first()
            
        if last_note:
            now = datetime.now(timezone.utc)
            
            # FIX: Convert naive datetime from DB (assumed UTC) to timezone-aware UTC
            last_activity_utc = last_note.updated_at.replace(tzinfo=timezone.utc)
            
            diff = now - last_activity_utc
            
            if diff.days > 0:
                last_activity = f"{diff.days}d ago"
            elif diff.seconds >= 3600:
                last_activity = f"{diff.seconds // 3600}h ago"
            elif diff.seconds >= 60:
                last_activity = f"{diff.seconds // 60}m ago"
            else:
                last_activity = "Just now"
        else:
            last_activity = "Never"
            
        return jsonify({
            'success': True,
            'total_notes': total_notes,
            'last_activity': last_activity
        })
    except Exception as e:
        logger.error(f"Error fetching profile stats: {e}")
        return jsonify({
            'success': False,
            'message': 'Failed to load statistics'
        }), 500

# ================= PROFILE UPDATE API =================
@app.route('/api/profile/update', methods=['PUT'])
@login_required
@csrf.exempt
@require_ajax_csrf
def update_profile():
    try:
        data = request.get_json()

        if not data:
            return jsonify({
                "success": False,
                "message": "No data received"
            }), 400

        user = current_user

        # Update fields safely
        user.first_name = data.get('first_name', '').strip()
        user.last_name = data.get('last_name', '').strip()
        user.mobile_number = data.get('mobile_number', '').strip()
        user.bio = data.get('bio', '').strip()

        db.session.commit()

        return jsonify({
            "success": True,
            "message": "Profile updated successfully!"
        }), 200

    except Exception as e:
        db.session.rollback()
        logger.error("Profile Update Error:", e)

        return jsonify({
            "success": False,
            "message": "Server error while updating profile"
        }), 500


@app.route('/api/profile/update-v2', methods=['POST'])
@login_required
@limiter.limit("10 per hour")
def api_update_profile():
    """API endpoint to update user profile information"""
    try:
        data = request.get_json()
        
        # Validate data exists
        if not data:
            return jsonify({
                'success': False,
                'message': 'No data provided'
            }), 400
        
        # Track if any changes were made
        changes_made = False
        
        # Update first name (with safe attribute check)
        if 'first_name' in data:
            new_first_name = data['first_name'].strip()
            old_first_name = getattr(current_user, 'first_name', None)
            if old_first_name != new_first_name:
                if hasattr(current_user, 'first_name'):
                    current_user.first_name = new_first_name
                else:
                    setattr(current_user, 'first_name', new_first_name)
                changes_made = True
        
        # Update last name (with safe attribute check)
        if 'last_name' in data:
            new_last_name = data['last_name'].strip()
            old_last_name = getattr(current_user, 'last_name', None)
            if old_last_name != new_last_name:
                if hasattr(current_user, 'last_name'):
                    current_user.last_name = new_last_name
                else:
                    setattr(current_user, 'last_name', new_last_name)
                changes_made = True
        
        # Update bio (with safe attribute check)
        if 'bio' in data:
            new_bio = data['bio'].strip()
            if len(new_bio) > 150:
                return jsonify({
                    'success': False,
                    'message': 'Bio must be 150 characters or less'
                }), 400
            old_bio = getattr(current_user, 'bio', None)
            if old_bio != new_bio:
                if hasattr(current_user, 'bio'):
                    current_user.bio = new_bio
                else:
                    setattr(current_user, 'bio', new_bio)
                changes_made = True
        
        # Only commit if changes were made
        if changes_made:
            db.session.commit()
            logger.info(f"Profile updated successfully for user {current_user.username}")
        else:
            logger.info(f"No changes detected for user {current_user.username}")
        
        # Return updated data (with safe attribute access)
        return jsonify({
            'success': True,
            'message': 'Profile updated successfully!',
            'data': {
                'first_name': getattr(current_user, 'first_name', '') or '',
                'last_name': getattr(current_user, 'last_name', '') or '',
                'bio': getattr(current_user, 'bio', '') or '',
                'profile_picture_url': getattr(current_user, 'profile_picture_url', None)
            }
        })
        
    except AttributeError as ae:
        db.session.rollback()
        logger.error(f"AttributeError in profile update: {ae}")
        logger.error(f"Missing columns in User model. Please run database migration.")
        return jsonify({
            'success': False,
            'message': 'Database schema needs update. Please contact administrator.'
        }), 500
    except Exception as e:
        db.session.rollback()
        logger.error(f"Error updating profile for user {current_user.username}: {e}")
        logger.error(f"Traceback: {traceback.format_exc()}")
        return jsonify({
            'success': False,
            'message': 'Failed to update profile. Please try again.'
        }), 500


import os
from werkzeug.utils import secure_filename
from flask import current_app

# Allowed image types
ALLOWED_EXTENSIONS = {'png', 'jpg', 'jpeg', 'gif', 'webp'}

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS


# ================= PROFILE PICTURE UPLOAD =================
@app.route('/api/profile/upload-picture', methods=['POST'])
@login_required
def upload_profile_picture():
    try:

        # 1️⃣ Check file exists
        if 'picture' not in request.files:
            return jsonify({
                "success": False,
                "message": "No file uploaded"
            }), 400

        file = request.files['picture']

        # 2️⃣ Empty filename
        if file.filename == '':
            return jsonify({
                "success": False,
                "message": "No selected file"
            }), 400

        # 3️⃣ Validate extension
        if not allowed_file(file.filename):
            return jsonify({
                "success": False,
                "message": "Invalid file type. Only images allowed."
            }), 400

        # 4️⃣ Secure filename
        filename = secure_filename(file.filename)

        # unique filename
        ext = filename.rsplit('.', 1)[1].lower()
        new_filename = f"user_{current_user.id}.{ext}"

        # 5️⃣ Save path
        upload_folder = os.path.join(current_app.root_path, 'static/profile_pics')
        os.makedirs(upload_folder, exist_ok=True)

        file_path = os.path.join(upload_folder, new_filename)
        file.save(file_path)

        # 6️⃣ Save in DB
        current_user.profile_picture_url = f"/static/profile_pics/{new_filename}"
        db.session.commit()

        return jsonify({
            "success": True,
            "message": "Profile picture updated!",
            "new_url": current_user.profile_picture_url
        })

    except Exception as e:
        logger.error("UPLOAD ERROR:", e)
        return jsonify({
            "success": False,
            "message": "Upload failed"
        }), 500



@app.route('/api/profile/remove-picture', methods=['POST'])
@login_required
@limiter.limit("10 per hour")
def api_remove_profile_picture():
    """API endpoint to remove user profile picture"""
    try:
        # Delete old file if exists
        old_picture_url = getattr(current_user, 'profile_picture_url', None)
        if old_picture_url and 'default_profile.png' not in old_picture_url:
            try:
                upload_folder = app.config.get('UPLOAD_FOLDER', 'static/profile_pics')
                user_folder = os.path.join(upload_folder, str(current_user.id))
                old_filename = old_picture_url.split('/')[-1]
                old_filepath = os.path.join(user_folder, old_filename)
                
                if os.path.exists(old_filepath):
                    os.remove(old_filepath)
                    logger.info(f"Deleted profile picture file: {old_filepath}")
            except Exception as e:
                logger.warning(f"Failed to delete old profile picture file: {e}")
        
        # Clear profile picture URL in database
        if hasattr(current_user, 'profile_picture_url'):
            current_user.profile_picture_url = None
        
        db.session.commit()
        
        logger.info(f"Profile picture removed for user {current_user.username}")
        
        return jsonify({
            'success': True,
            'message': 'Profile picture removed successfully!'
        })
        
    except Exception as e:
        db.session.rollback()
        logger.error(f"Error removing profile picture: {e}")
        return jsonify({
            'success': False,
            'message': 'Failed to remove profile picture'
        }), 500



@app.route('/api/account/delete', methods=['POST'])
@login_required
def api_delete_account():
    """Permanently delete account, notes and sessions in one safe transaction."""
    try:
        data = request.get_json(silent=True) or {}
        confirm_username = data.get('confirm_username', '').strip()

        if confirm_username != current_user.username:
            return jsonify({
                'success': False,
                'message': 'Username confirmation does not match.'
            }), 400

        user_id = current_user.id

        # Save email + mobile BEFORE logout/delete for blocking re-registration
        deleted_email  = current_user.email
        deleted_mobile = getattr(current_user, 'mobile_number', None)
        deleted_uname  = current_user.username

        # Logout FIRST so session is gone before DB changes
        session.clear()
        logout_user()

        # Delete in proper dependency order
        UserSession.query.filter_by(user_id=user_id).delete()
        Note.query.filter_by(user_id=user_id).delete()
        Collaborator.query.filter_by(collaborator_id=user_id).delete()

        user = User.query.get(user_id)
        if user:
            db.session.delete(user)

        # Record deleted account so email/mobile cannot be reused
        deleted_record = DeletedAccount(
            email         = deleted_email,
            mobile_number = deleted_mobile,
            username      = deleted_uname,
        )
        db.session.add(deleted_record)

        db.session.commit()
        logger.warning(f"Account permanently deleted: user_id={user_id}")

        return jsonify({
            'success': True,
            'message': 'Account deleted successfully.'
        }), 200

    except Exception as e:
        db.session.rollback()
        logger.error(f"Error deleting account: {e}")
        return jsonify({
            'success': False,
            'message': 'Failed to delete account. Please try again.'
        }), 500



# Add this route to your app.py file (around line 850, after the download_pdf route)

@app.route('/export_data')
@login_required
@limiter.limit("5 per hour")
@limiter.limit("20 per day")
def export_data():
    """
    Export all user notes as JSON.
    Uses chunked DB query (yield_per) to avoid loading all notes into RAM at once.
    Safe for users with thousands of notes.
    """
    try:
        # Stats queries — lightweight aggregates only
        total_notes    = Note.query.filter_by(user_id=current_user.id).count()
        private_count  = Note.query.filter_by(user_id=current_user.id, is_private=True).count()
        favorite_count = Note.query.filter_by(user_id=current_user.id, is_favorite=True).count()

        # Category breakdown — single DB round-trip
        from sqlalchemy import func
        cat_rows = (
            db.session.query(Note.category, func.count(Note.id))
            .filter_by(user_id=current_user.id)
            .group_by(Note.category)
            .all()
        )
        category_counts = {(cat or 'General'): cnt for cat, cnt in cat_rows}

        # Build JSON incrementally into BytesIO (no full list in RAM)
        buffer = BytesIO()

        header = json.dumps({
            'user': {
                'username':    current_user.username,
                'email':       current_user.email,
                'is_premium':  current_user.is_premium,
                'export_date': datetime.now(timezone.utc).isoformat(),
            },
            'statistics': {
                'total_notes':    total_notes,
                'private_notes':  private_count,
                'favorite_notes': favorite_count,
                'categories':     category_counts,
            },
        }, ensure_ascii=False, indent=2)

        # Write header + open notes array
        buffer.write(header[:-1].encode('utf-8'))   # everything except closing }
        buffer.write(b',\n  "notes": [\n')

        # Stream notes 50 at a time — yield_per avoids full load
        first = True
        query = (
            Note.query
            .filter_by(user_id=current_user.id)
            .order_by(Note.id)
            .yield_per(50)
        )
        for note in query:
            try:
                soup = BeautifulSoup(note.content or '', 'html.parser')
                plain = soup.get_text(separator='\n', strip=True)
            except Exception:
                plain = ''

            note_obj = {
                'id':            note.id,
                'title':         note.title,
                'content_html':  note.content,
                'content_plain': plain,
                'category':      note.category or 'General',
                'paper_size':    note.paper_size,
                'is_favorite':   note.is_favorite,
                'is_private':    note.is_private,
                'created_at':    note.created_at.isoformat(),
                'updated_at':    note.updated_at.isoformat(),
            }
            chunk = ('    ' if first else ',\n    ') + json.dumps(note_obj, ensure_ascii=False)
            buffer.write(chunk.encode('utf-8'))
            first = False

        # Close notes array + root object
        buffer.write(b'\n  ]\n}')
        buffer.seek(0)

        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        filename  = f"notesaver_export_{current_user.username}_{timestamp}.json"

        logger.info(f"Data exported by {current_user.username}: {total_notes} notes")
        return send_file(
            buffer,
            as_attachment=True,
            download_name=filename,
            mimetype='application/json'
        )

    except Exception as e:
        logger.error(f"Error exporting data for {current_user.username}: {e}")
        flash('Error exporting data. Please try again.', 'error')
        return redirect(url_for('settings'))






# ===== SETTINGS API ROUTES =====

@app.route('/api/password/change', methods=['POST'])
@login_required
@limiter.limit("5 per hour")
def api_change_password():
    """API endpoint for changing user password"""
    try:
        data = request.get_json()
        current_password = data.get('current_password')
        new_password = data.get('new_password')
        
        if not current_password or not new_password:
            return jsonify({
                'success': False,
                'message': 'Both current and new passwords are required.'
            }), 400
        
        # Verify current password
        if not current_user.check_password(current_password):
            logger.warning(f"Password change failed for user {current_user.username}: Incorrect current password")
            return jsonify({
                'success': False,
                'message': 'Current password is incorrect.'
            }), 400
        
        # Validate new password strength
        if len(new_password) < 8:
            return jsonify({
                'success': False,
                'message': 'New password must be at least 8 characters long.'
            }), 400

        if new_password == current_password:
            return jsonify({
                'success': False,
                'message': 'New password must be different from current password.'
            }), 400

        # Require at least 1 digit and 1 letter
        if not any(c.isdigit() for c in new_password) or not any(c.isalpha() for c in new_password):
            return jsonify({
                'success': False,
                'message': 'Password must contain at least one letter and one number.'
            }), 400
        
        # Update password
        current_user.set_password(new_password)
        db.session.commit()
        
        logger.info(f"Password changed successfully for user {current_user.username}")
        
        return jsonify({
            'success': True,
            'message': 'Password updated successfully!'
        })
        
    except Exception as e:
        db.session.rollback()
        logger.error(f"Error changing password for user {current_user.username}: {e}")
        return jsonify({
            'success': False,
            'message': 'An error occurred. Please try again.'
        }), 500


@app.route('/api/preferences/update', methods=['POST'])
@login_required
@limiter.limit("20 per hour")
def api_update_preferences():
    """API endpoint for updating user preferences"""
    try:
        data = request.get_json()
        app_theme = data.get('appTheme', 'light')
        default_view = data.get('defaultView', 'list')
        
        # Store preferences in session (you can also add columns to User model)
        session['app_theme'] = app_theme
        session['default_view'] = default_view
        
        logger.info(f"Preferences updated for user {current_user.username}: theme={app_theme}, view={default_view}")
        
        return jsonify({
            'success': True,
            'message': 'Preferences saved successfully!'
        })
        
    except Exception as e:
        logger.error(f"Error updating preferences for user {current_user.username}: {e}")
        return jsonify({
            'success': False,
            'message': 'Failed to save preferences.'
        }), 500


@app.route('/api/notes/delete_all', methods=['POST'])
@login_required
@limiter.limit("3 per day")
def api_delete_all_notes():
    """
    Delete all OR selected notes.
    Body (optional): {"note_ids": [1, 2, 3]}  → deletes only those
    No body → deletes ALL user notes.
    """
    try:
        data = request.get_json(silent=True) or {}
        note_ids = data.get('note_ids')

        query = Note.query.filter_by(user_id=current_user.id)

        if note_ids:
            if not isinstance(note_ids, list) or not all(isinstance(i, int) for i in note_ids):
                return jsonify({'success': False, 'message': 'note_ids must be a list of integers.'}), 400
            query = query.filter(Note.id.in_(note_ids))

        note_count = query.count()
        if note_count == 0:
            return jsonify({'success': False, 'message': 'No matching notes found.'}), 400

        query.delete(synchronize_session='fetch')
        db.session.commit()

        action = f"bulk-selected {len(note_ids)}" if note_ids else "all"
        logger.warning(f"[{action}] notes deleted by {current_user.username}: {note_count} notes")

        return jsonify({
            'success': True,
            'message': f'Successfully deleted {note_count} note(s).',
            'deleted_count': note_count,
        })

    except Exception as e:
        db.session.rollback()
        logger.error(f"Error in bulk delete by {current_user.username}: {e}")
        return jsonify({'success': False, 'message': 'Failed to delete notes.'}), 500


@app.route('/api/sessions')
@login_required
def api_sessions():

    current_token = session.get("session_token")

    user_sessions = UserSession.query.filter_by(
        user_id=current_user.id,
        is_active=True
    ).order_by(UserSession.last_activity.desc()).all()

    data = []

    now = datetime.utcnow()

    for s in user_sessions:

        is_current = (s.session_token == current_token)

        # Human-readable last activity — timezone-aware fix
        last_activity_utc = s.last_activity
        if last_activity_utc.tzinfo is not None:
            # aware → naive UTC
            from datetime import timezone as _tz
            last_activity_utc = last_activity_utc.astimezone(_tz.utc).replace(tzinfo=None)
        diff = now - last_activity_utc
        secs = max(0, int(diff.total_seconds()))
        if secs < 60:
            last_activity = "Just now"
        elif secs < 3600:
            last_activity = f"{secs // 60}m ago"
        elif secs < 86400:
            last_activity = f"{secs // 3600}h ago"
        else:
            last_activity = s.last_activity.strftime("%d %b %Y, %I:%M %p")

        data.append({
            "id": s.session_token,
            "device": s.device or "Unknown Device",
            "location": s.location or "Detecting...",
            "last_activity": last_activity,
            "ip": s.ip_address or "",
            "status": "Current" if is_current else "Active",
            "is_current": is_current
        })

    return jsonify(data)











@app.route("/api/cleanup-old-sessions", methods=["POST"])
@login_required
def cleanup_old_sessions():
    MAX_ACTIVE_SESSIONS = 2
    active_sessions = UserSession.query.filter_by(
        user_id=current_user.id,
        is_active=True
    ).order_by(UserSession.created_at.asc()).all()

    current_token = session.get("session_token")
    removed = 0

    sessions_except_current = [s for s in active_sessions if s.session_token != current_token]
    excess = len(active_sessions) - MAX_ACTIVE_SESSIONS

    if excess > 0:
        to_remove = sessions_except_current[:excess]
        for s in to_remove:
            s.is_active = False
            removed += 1
        db.session.commit()
        logger.info(f"Cleaned up {removed} excess sessions for {current_user.username}")

    return jsonify({"success": True, "removed": removed, "active_now": len(active_sessions) - removed})


@app.route("/api/logout-session/<token>", methods=["POST"])
@login_required
@csrf.exempt
@require_ajax_csrf
def logout_session(token):

    # current session logout allowed nahi — sirf doosre ka karo
    if token == session.get("session_token"):
        return jsonify({"success": False, "error": "Cannot logout current session"}), 400

    # sirf is user ki session, aur sirf active wali
    user_session = UserSession.query.filter_by(
        user_id=current_user.id,
        session_token=token,
        is_active=True
    ).first()

    if not user_session:
        # Already logged out — silently succeed so UI refreshes cleanly
        return jsonify({"success": True, "message": "Session already inactive"})

    user_session.is_active = False
    user_session.last_activity = datetime.now(timezone.utc)
    db.session.commit()

    logger.info(f"Session logged out by user {current_user.username}: {token[:8]}...")
    return jsonify({"success": True})


@app.route('/api/sessions/logout_all_others', methods=['POST'])
@login_required
def logout_all_others():
    current_token = session.get("session_token")
    UserSession.query.filter(
        UserSession.user_id == current_user.id,
        UserSession.session_token != current_token
    ).update({"is_active": False})
    db.session.commit()
    return jsonify({"success": True, "message": "All other sessions logged out"})


@app.route('/api/sessions/logout_all', methods=['POST'])
@login_required
@csrf.exempt
@require_ajax_csrf
def logout_all_sessions():
    # Sabke sabhi sessions logout karo including current - force logout
    UserSession.query.filter_by(
        user_id=current_user.id
    ).update({"is_active": False})
    db.session.commit()
    session.clear()
    logout_user()
    logger.info(f"All sessions force-logged out (self-initiated)")
    return jsonify({"success": True, "message": "All sessions logged out"})





# Add these routes to your app.py file

@app.route('/api/reset-note-password', methods=['POST'])
@limiter.limit("3 per hour")
@limiter.limit("10 per day")
def api_reset_note_password():
    """API endpoint for requesting note password reset"""
    try:
        # Enhanced logging
        logger.info(f"Reset request received from: {get_remote_address()}")
        logger.info(f"Request method: {request.method}")
        logger.info(f"Content-Type: {request.headers.get('Content-Type')}")
        
        # Check if request has JSON data
        if not request.is_json:
            logger.warning("Request is not JSON")
            return jsonify({
                'success': False, 
                'message': 'Invalid request format. JSON required.'
            }), 400
        
        # Get JSON data
        data = request.get_json()
        logger.info(f"Parsed JSON data: {data}")
        
        # Validate email field exists
        if not data or 'email' not in data:
            logger.warning("Email field missing in request")
            return jsonify({
                'success': False, 
                'message': 'Email address is required.'
            }), 400
        
        email = data.get('email', '').strip().lower()
        logger.info(f"Email extracted: {email}")
        
        # Validate email format
        email_regex = r'^[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}$'
        if not email or not re.match(email_regex, email):
            logger.warning(f"Invalid email format: {email}")
            return jsonify({
                'success': False, 
                'message': 'Please enter a valid email address.'
            }), 400
        
        # Find user by email
        user = User.query.filter_by(email=email).first()
        
        # SECURITY: Always return success message (don't reveal if email exists)
        if user:
            try:
                # Send reset email
                send_note_password_reset_email(user)
                logger.info(f"✅ Note password reset email sent to: {email}")
            except Exception as email_error:
                # Log error but still return success to user
                logger.error(f"❌ Failed to send email to {email}: {email_error}")
                logger.error(f"Email error details: {str(email_error)}")
                # Check if it's a mail configuration issue
                if 'MAIL_SERVER' not in app.config or not app.config.get('MAIL_SERVER'):
                    logger.error("⚠️ MAIL_SERVER not configured!")
                    return jsonify({
                        'success': False,
                        'message': 'Email service is not configured. Please contact administrator.'
                    }), 500
        else:
            logger.warning(f"⚠️ Note password reset requested for non-existent email: {email}")
        
        # Always return success (security best practice)
        return jsonify({
            'success': True,
            'message': 'If your email exists in our system, you will receive password reset instructions shortly.'
        }), 200
        
    except Exception as e:
        logger.error(f"❌ Critical error in api_reset_note_password: {e}")
        logger.error(f"Error type: {type(e).__name__}")
        logger.error(f"Error details: {str(e)}")
        logger.error(f"Traceback: {traceback.format_exc()}")
        
        return jsonify({
            'success': False,
            'message': 'An unexpected error occurred. Please try again later.'
        }), 500


# ALSO FIX: send_note_password_reset_email function
def send_note_password_reset_email(user):
    """Send note password reset email - MUST BE DEFINED BEFORE ROUTES"""
    try:
        # Check if mail is configured
        if not app.config.get('MAIL_SERVER'):
            logger.error("MAIL_SERVER not configured in app.config")
            raise Exception("Email service not configured")
        
        # Generate token
        token = serializer.dumps(user.email, salt='note-password-reset-salt')
        reset_url = url_for('selective_note_reset', token=token, _external=True)

        # ✅ FIX: Count notes with is_private=True (includes both proper password-protected and edge cases)
        protected_count = Note.query.filter(
            Note.user_id == user.id,
            Note.is_private == True
        ).count()
        
        if protected_count == 0:
            logger.info(f"No protected notes found for user {user.username}")
        
        # Create message
        msg = Message(
            "Note Password Reset Request - NoteSaver Pro",
            sender=app.config.get('MAIL_DEFAULT_SENDER', 'noreply@notesaver.com'),
            recipients=[user.email]
        )
        
        # Plain text version
        msg.body = f"""Hello {user.username},

You requested to reset passwords for your protected notes.

You currently have {protected_count} password-protected note(s).

To choose which notes to reset, click the link below:
{reset_url}

If you did not request this reset, please ignore this email.

This link will expire in 1 hour for security reasons.

Best regards,
NoteSaver Pro Team"""
        
        # HTML version
        msg.html = f"""
        <div style="font-family: Arial, sans-serif; max-width: 600px; margin: 0 auto;">
            <div style="background: #f8f9fa; padding: 20px; border-radius: 8px;">
                <h2 style="color: #dc3545;">🔐 Note Password Reset</h2>
                <p>Hello <strong>{user.username}</strong>,</p>
                <p>You requested to reset passwords for your protected notes.</p>
                <div style="background: #fff3cd; border: 1px solid #ffeaa7; padding: 15px; border-radius: 5px; margin: 20px 0;">
                    <p><strong>📊 Status:</strong> You have <strong>{protected_count}</strong> password-protected note(s).</p>
                </div>
                <div style="text-align: center; margin: 30px 0;">
                    <a href="{reset_url}" 
                       style="background: #dc3545; color: white; padding: 12px 25px; text-decoration: none; border-radius: 5px; display: inline-block;">
                        Select Notes to Reset
                    </a>
                </div>
                <hr style="margin: 30px 0;">
                <p><small>If you did not request this, ignore this email.</small></p>
                <p><small>Link expires in 1 hour.</small></p>
            </div>
        </div>
        """
        
        # Send email
        mail.send(msg)
        logger.info(f"✅ Note password reset email sent to: {user.email}")
        
    except Exception as e:
        logger.error(f"❌ Error sending note password reset email: {e}")
        logger.error(f"Traceback: {traceback.format_exc()}")
        raise


@app.route('/api/protected-notes')
@login_required
@limiter.limit("20 per minute")
def api_protected_notes():
    """API endpoint to get user's protected notes for selection"""
    try:
        # ✅ FIX: sirf password hash wale notes — is_private ek alag flag hai
        protected_notes = (
            Note.query
            .filter(
                Note.user_id == current_user.id,
                Note.private_password_hash.isnot(None)
            )
            .with_entities(Note.id, Note.title, Note.category, Note.created_at, Note.updated_at)
            .order_by(Note.updated_at.desc())
            .limit(500)
            .all()
        )
        
        notes_data = [{
            'id': note.id,
            'title': note.title,
            'category': note.category,
            'created_at': note.created_at.strftime('%d %b %Y'),
            'updated_at': note.updated_at.strftime('%d %b %Y')
        } for note in protected_notes]
        
        logger.info(f"Protected notes API accessed by user {current_user.username}: {len(notes_data)} notes")
        
        return jsonify({
            'success': True,
            'notes': notes_data,
            'count': len(notes_data)
        })
        
    except Exception as e:
        logger.error(f"Error fetching protected notes for user {current_user.username}: {e}")
        return jsonify({
            'success': False,
            'message': 'Error loading notes'
        }), 500
        
@app.route('/selective-note-reset/<token>', methods=['GET', 'POST'])
@limiter.limit("5 per hour")
def selective_note_reset(token):
    try:
        email = serializer.loads(token, salt='note-password-reset-salt', max_age=3600)
    except SignatureExpired:
        flash('The reset link has expired.', 'danger')
        return redirect(url_for('login'))
    except BadSignature:
        flash('Invalid reset token.', 'danger')
        return redirect(url_for('login'))

    user = User.query.filter_by(email=email).first_or_404()

    if request.method == 'GET':
        # ✅ FIX: Fetch both notes with password hash AND notes with is_private=True
        # (handles both proper password-protected notes and edge cases from bugs)
        protected_notes = Note.query.filter(
            Note.user_id == user.id,
            Note.is_private == True
        ).order_by(Note.updated_at.desc()).all()

        # Auto-heal: ensure is_private notes actually have password hashes
        # If is_private=True but no hash, this is a recovery case
        healed = False
        for note in protected_notes:
            if not note.private_password_hash:
                # Generate a default hash for this note to make it selectable
                # User will be able to remove protection from it
                note.private_password_hash = generate_password_hash(
                    f"auto-recovered-{note.id}-{user.id}"
                )
                healed = True
        
        if healed:
            try:
                db.session.commit()
                logger.info(f"Auto-healed password hashes for {user.username}'s private notes")
            except Exception:
                db.session.rollback()

        return render_template('choose_notes_to_reset.html', user=user, token=token, protected_notes=protected_notes)

    elif request.method == 'POST':
        selected_ids = request.form.getlist('selected_notes')
        if not selected_ids:
            flash("No notes selected for reset.", "warning")
            return redirect(request.url)

        try:
            # ✅ FIX: Match on id+user_id only — don't require is_private=True
            # (handles the inconsistent-state notes from old code paths)
            notes = Note.query.filter(
                Note.id.in_(selected_ids),
                Note.user_id == user.id
            ).all()

            for note in notes:
                note.is_private = False
                note.private_password_hash = None

            db.session.commit()
            flash(f"Password protection removed from {len(notes)} notes.", "success")
            return redirect(url_for('login'))

        except Exception as e:
            db.session.rollback()
            logger.error(f"Selective reset error: {e}")
            flash("An error occurred while resetting selected notes.", "danger")
            return redirect(url_for('login'))

# ------------------- HEALTH CHECK & MONITORING ROUTES ---------------------

@app.route('/health')
@limiter.exempt  # Health check ko exempt rakhte hain
def health_check():
    """Health check endpoint for monitoring"""
    return jsonify({
        "status": "healthy",
        "timestamp": datetime.utcnow().isoformat(),
        "redis_available": REDIS_AVAILABLE,
        "version": "1.0.0"
    })

@app.route('/api/rate-limit-status')
@login_required
@limiter.limit("10 per minute")
def rate_limit_status():
    """Check current rate limit status for user"""
    try:
        user_tier = get_user_tier()
        user_id = get_user_id()
        
        # Basic rate limit info
        status = {
            "user_tier": user_tier,
            "user_id": user_id if user_tier != "anonymous" else "anonymous",
            "rate_limits": {
                "login": "10/minute, 50/hour",
                "create_note": "30/minute, 200/hour",
                "edit_note": "50/minute",
                "pdf_download": "10/minute, 50/hour" if user_tier == "premium" else "Not available"
            }
        }
        
        # Add Redis-based PDF download info if available
        if REDIS_AVAILABLE and user_tier in ["authenticated", "premium"]:
            try:
                pdf_download_key = f"pdf_downloads_{current_user.id}"
                daily_downloads = redis_client.get(pdf_download_key) or 0
                max_daily_pdfs = 50 if current_user.is_premium_active() else 10
                
                status["pdf_downloads"] = {
                    "used_today": int(daily_downloads),
                    "limit": max_daily_pdfs,
                    "remaining": max_daily_pdfs - int(daily_downloads)
                }
            except Exception as e:
                logger.error(f"Error fetching PDF download status: {e}")
        
        return jsonify(status)
        
    except Exception as e:
        logger.error(f"Error in rate_limit_status: {e}")
        return jsonify({"error": "Unable to fetch status"}), 500

# ------------------- ENHANCED SEARCH ROUTE ---------------------

@app.route('/api/notes/search')
@login_required
@limiter.limit("60 per minute")
def api_search_notes():
    """Fast JSON search API — usable by JS without page reload."""
    q = request.args.get('q', '').strip()
    category = request.args.get('category', '').strip()
    page = request.args.get('page', 1, type=int)

    if not q and not category:
        return jsonify({'success': False, 'message': 'Search query required.'}), 400

    query = Note.query.filter_by(user_id=current_user.id)
    if q:
        pattern = f"%{q}%"
        query = query.filter(
            (Note.title.ilike(pattern)) | (Note.content.ilike(pattern))
        )
    if category:
        query = query.filter_by(category=category)

    paginated = query.order_by(Note.updated_at.desc()).paginate(
        page=page, per_page=10, error_out=False
    )

    results = [{
        'id': n.id,
        'title': n.title,
        'category': n.category,
        'preview': get_note_preview(n.content, 120),
        'updated_at': n.updated_at.strftime('%d %b %Y'),
        'is_private': n.is_private,
        'is_favorite': n.is_favorite,
    } for n in paginated.items]

    return jsonify({
        'success': True,
        'results': results,
        'total': paginated.total,
        'pages': paginated.pages,
        'page': page,
    })


# ─── Existing page-based search ─────────────────────────────────────────────
@app.route('/search')
@login_required
def search():
    """Enhanced search with rate limiting based on user tier"""
    user_tier = get_user_tier()
    
    # Apply different search limits based on user tier
    if user_tier == "premium":
        @limiter.limit("100 per minute")
        def premium_search():
            return perform_search()
        return premium_search()
    elif user_tier == "authenticated":
        @limiter.limit("50 per minute")
        def auth_search():
            return perform_search()
        return auth_search()
    else:
        @limiter.limit("10 per minute")
        def anon_search():
            return perform_search()
        return anon_search()

def perform_search():
    """Actual search logic"""
    query = request.args.get('q', '')
    if not query:
        return redirect(url_for('dashboard'))
    
    # Perform search logic here
    search_pattern = f"%{query}%"
    notes = Note.query.filter_by(user_id=current_user.id).filter(
        (Note.title.ilike(search_pattern)) | (Note.content.ilike(search_pattern))
    ).order_by(Note.updated_at.desc()).limit(20).all()
    
    logger.info(f"Search performed by user {current_user.username}: '{query}' - {len(notes)} results")
    
    return render_template('search_results.html', notes=notes, query=query)

# ------------------- Error Handlers ---------------------

@app.errorhandler(404)
def page_not_found(e):
    logger.warning(f"404 error: {request.url}")
    return render_template('404.html'), 404

@app.errorhandler(403)
def forbidden(e):
    logger.warning(f"403 error: {request.url} by {get_remote_address()}")
    return render_template('403.html'), 403

@app.errorhandler(500)
def internal_error(e):
    db.session.rollback()
    logger.error(f"500 error: {str(e)} at {request.url}")
    return render_template('500.html'), 500

@app.errorhandler(413)
def request_entity_too_large(e):
    """Handle file upload size limit exceeded"""
    logger.warning(f"413 error - File too large from {get_remote_address()}")
    flash('File size too large. Please upload a smaller file.', 'error')
    return redirect(url_for('dashboard')), 413

# ------------------- ADMIN ROUTES (Optional) ---------------------

@app.route('/admin/stats')
@login_required
@limiter.limit("10 per minute")
def admin_stats():
    """Admin statistics - only if user is admin"""
    if not getattr(current_user, 'is_admin', False):
        flash('Access denied. Admin privileges required.', 'danger')
        return redirect(url_for('dashboard'))
    
    try:
        stats = {
            "total_users": User.query.count(),
            "total_notes": Note.query.count(),
            "premium_users": User.query.filter_by(is_premium=True).count(),
            "private_notes": Note.query.filter_by(is_private=True).count(),
            "redis_status": REDIS_AVAILABLE
        }
        
        logger.info(f"Admin stats accessed by user {current_user.username}")
        return jsonify(stats)
        
    except Exception as e:
        logger.error(f"Error in admin_stats: {e}")
        return jsonify({"error": "Unable to fetch stats"}), 500

# ------------------- RATE LIMIT MONITORING ---------------------

@app.before_request
def before_request():
    import time
    request.start_time = time.time()

    if request.endpoint in ['login', 'register', 'verify_note_password', 'forgot_password']:
        logger.info(f"Security endpoint accessed: {request.endpoint} from {get_remote_address()}")
@app.before_request
def update_last_activity():
    # Static files pe check mat karo — performance waste
    if request.endpoint and request.endpoint.startswith('static'):
        return

    token = session.get("session_token")
    if not token:
        return

    # ── CRITICAL CHECK: Kya yeh session DB mein active hai? ──────────
    # Agar dusre device ne login kiya aur is session ko deactivate kar diya
    # toh is browser ko bhi force logout karo
    s = UserSession.query.filter_by(session_token=token).first()

    if not s or not s.is_active:
        # Session DB mein inactive hai — force logout karo
        if current_user.is_authenticated:
            logger.info(
                f"Force logout: session {token[:8]} is inactive in DB "
                f"for user {current_user.username} — logging out"
            )
        session.clear()
        logout_user()
        # Sirf login/logout routes pe redirect mat karo
        if request.endpoint not in ('login', 'logout', 'index', 'static'):
            from flask import abort
            return redirect(url_for('login'))
        return

    # Session active hai — last activity update karo
    s.last_activity = datetime.now(timezone.utc)
    db.session.commit()

@app.after_request
def after_request(response):
    # ── Security headers ──────────────────────────────────
    response.headers['X-Content-Type-Options'] = 'nosniff'
    response.headers['X-Frame-Options'] = 'DENY'
    response.headers['X-XSS-Protection'] = '1; mode=block'
    response.headers['Referrer-Policy'] = 'strict-origin-when-cross-origin'
    response.headers['Permissions-Policy'] = 'geolocation=(self), camera=(), microphone=()'

    # Only add HSTS on HTTPS
    if request.is_secure:
        response.headers['Strict-Transport-Security'] = 'max-age=31536000; includeSubDomains'

    # ── Slow request logging ──────────────────────────────
    if hasattr(request, 'start_time'):
        duration = time.time() - request.start_time
        if duration > 2.0:
            logger.warning(f"⚠️ Slow request: {request.endpoint} took {duration:.2f}s")

    return response

# ------------------- App Initialization ---------------------



if __name__ == '__main__':
    # Production mein debug=False karna chahiye
    logger.info("Starting Flask application with enhanced rate limiting")
    logger.info(f"Redis available: {REDIS_AVAILABLE}")
    app.run(debug=True)
